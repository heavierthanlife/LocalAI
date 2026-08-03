"""Batch audit orchestrator — pure business logic extracted from routes.

Provides:
  - RiskScorer:     configurable risk formula (snapshot-locked via W1a tests)
  - SUB_CHECKERS:   registry of non-blocking sub-checkers
  - DOCXBuilder:    shared DOCX report builder (replaces ExcelBuilder)
  - HTMLBuilder:    builds the full comparison report HTML
"""
import os
import html
import json
import logging
from datetime import datetime, timezone
from io import BytesIO

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from openpyxl.utils import get_column_letter

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.services.batch_compare_svc import (
    _precompute_tfidf_for_files,
    _compute_pair_similarity_from_matrix,
)
from app.services.file_processing import (
    preprocess_text_for_similarity,
    keyword_overlap_similarity,
    extract_keywords,
    compute_similarity_with_numbers,
    remove_template_content,
    truncate_filename,
)

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════
# Risk scoring
# ═══════════════════════════════════════════════════════════════════════════

class RiskScorer:
    """Configurable risk scoring formula.

    The default weights match the current production formula
    (0.3*key + 0.3*attr + 0.2*text + 0.2*img), locked via
    tests/test_batch_orchestrator.py::test_snapshot_risk_formula.
    """

    WEIGHTS = {
        "key_info": 0.3,
        "file_attr": 0.3,
        "text_sim": 0.2,
        "image_sim": 0.2,
    }

    @classmethod
    def compute(cls, key_info_pct: float, file_attr_val: float,
                text_sim_pct: float, img_sim_val: float) -> float:
        return (
            cls.WEIGHTS["key_info"] * key_info_pct +
            cls.WEIGHTS["file_attr"] * file_attr_val +
            cls.WEIGHTS["text_sim"] * text_sim_pct +
            cls.WEIGHTS["image_sim"] * img_sim_val
        )


# ═══════════════════════════════════════════════════════════════════════════
# Pairwise comparison core
# ═══════════════════════════════════════════════════════════════════════════

def compute_single_pair(file_data, i, j, check_items, tfidf_matrix=None,
                        template_text=None):
    """Compute a single file-pair's similarity metrics."""
    text1 = file_data[i]['text']
    text2 = file_data[j]['text']
    meta1 = file_data[i]['metadata']
    meta2 = file_data[j]['metadata']
    images1 = file_data[i]['images']
    images2 = file_data[j]['images']

    from app.services.file_processing import image_similarity, file_attr_similarity

    # Image similarity
    img_sim = image_similarity(images1, images2) if check_items.get('image_sim', True) else 0.0

    # Text similarity (TF-IDF cosine)
    if check_items.get('text_sim', True) and tfidf_matrix is not None:
        sim = _compute_pair_similarity_from_matrix(tfidf_matrix, i, j)
    else:
        sim = 0.0

    # Key info overlap
    if check_items.get('key_info', True):
        t1 = preprocess_text_for_similarity(text1)
        t2 = preprocess_text_for_similarity(text2)
        if template_text:
            t1 = remove_template_content(t1, template_text)
            t2 = remove_template_content(t2, template_text)
        key_sim = keyword_overlap_similarity(t1, t2)
    else:
        key_sim = 0.0

    # File attribute similarity
    if check_items.get('file_attr', True) and meta1 and meta2:
        attr_sim = file_attr_similarity(meta1, meta2)
    else:
        attr_sim = 0.0

    text_sim_val = sim * 100
    key_info_val = key_sim * 100
    file_attr_val = attr_sim
    img_sim_val = img_sim

    risk = RiskScorer.compute(key_info_val, file_attr_val, text_sim_val, img_sim_val)

    _, html1, html2, blocks = compute_similarity_with_numbers(text1, text2, template_text)

    return {
        'i': i, 'j': j,
        'name1': file_data[i]['filename'],
        'name2': file_data[j]['filename'],
        'text1': text1, 'text2': text2,
        'sim': sim * 100,
        'risk': risk,
        'blocks': blocks,
        'html1': html1, 'html2': html2,
        'used_weights': {},
        'attr_same': 1 if meta1.get('author') and meta1['author'] == meta2.get('author') else 0,
    }


def compute_all_pairs(file_data, check_items, tfidf_matrix=None, template_text=None):
    """Run pairwise comparison for all file pairs."""
    n = len(file_data)
    pairs = []
    risk_matrix = [[0] * n for _ in range(n)]
    for i in range(n):
        for j in range(i + 1, n):
            pair = compute_single_pair(file_data, i, j, check_items, tfidf_matrix, template_text)
            pairs.append(pair)
            risk_matrix[i][j] = pair['risk']
            risk_matrix[j][i] = pair['risk']
    return pairs, risk_matrix


def build_key_info_matches(pairs):
    """Post-process key info matches from pairs."""
    matches = []
    for p in pairs:
        kw1 = set(extract_keywords(p['text1'], 20))
        kw2 = set(extract_keywords(p['text2'], 20))
        matches.append({
            'name1': p['name1'],
            'name2': p['name2'],
            'common_keywords': list(kw1 & kw2)[:10],
        })
    return matches


def build_attr_details(file_data):
    """Post-process file attribute details."""
    details = []
    for fd in file_data:
        meta = fd['metadata']
        details.append({
            'filename': fd['filename'],
            'author': meta.get('author', ''),
            'creation_date': meta.get('creationDate', ''),
            'creator': meta.get('creator', ''),
            'producer': meta.get('producer', ''),
        })
    return details


# ═══════════════════════════════════════════════════════════════════════════
# Sub-checker registry
# ═══════════════════════════════════════════════════════════════════════════

def _run_typo_check(file_data, user_id, thread_id, project_id=None, audit=None):
    try:
        from app.services.typo_detector import detect_typos_batch, save_typo_results as save_typo
        results = detect_typos_batch(file_data, audit=audit)
        save_typo(user_id, thread_id, results)
        return results
    except Exception as e:
        logger.warning(f"Typo detection failed (non-blocking): {e}")
        return None


def _run_relationship_check(file_data, user_id, thread_id, project_id=None, audit=None):
    try:
        from app.services.relationship_extractor import (
            extract_relationships as run_rel,
            save_relationship_results as save_rel,
        )
        report = run_rel(file_data, audit=audit)
        save_rel(user_id, thread_id, report, project_id=project_id)
        return report
    except Exception as e:
        logger.warning(f"Relationship extraction failed (non-blocking): {e}")
        return None


def _run_quote_check(file_data, user_id, thread_id, project_id=None, audit=None):
    try:
        from app.services.quote_anomaly import (
            compare_bidders_quotes,
            save_quote_anomaly_results,
        )
        result = compare_bidders_quotes(file_data, audit=audit)
        save_quote_anomaly_results(user_id, thread_id, result['per_bidder'], result, project_id=project_id)
        return result
    except Exception as e:
        logger.warning(f"Quote anomaly check failed (non-blocking): {e}")
        return None


# Registry: add a new checker = one line here + one HTML section method
SUB_CHECKERS = {
    'typo': _run_typo_check,
    'relationship': _run_relationship_check,
    'quote': _run_quote_check,
}


def run_all_sub_checkers(file_data, user_id, thread_id, project_id=None, audit=None):
    results = {}
    for name, fn in SUB_CHECKERS.items():
        try:
            results[name] = fn(file_data, user_id, thread_id, project_id=project_id, audit=audit)
        except Exception as e:
            logger.warning(f"Sub-checker {name} raised: {e}")
            results[name] = None
    return results


# ═══════════════════════════════════════════════════════════════════════════
# Excel workbook builder (shared)
# ═══════════════════════════════════════════════════════════════════════════

def build_excel_workbook(file_data, pairs, key_info_matches):
    """Build Excel workbook for batch comparison results.

    Shared by compare_batch() route and export_batch_docx_download().
    Eliminates the 60-line-duplicate Excel generation.
    """
    n = len(file_data)
    _ts = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')
    max_risk = max(p['risk'] for p in pairs) if pairs else 0
    max_sim = max(p['sim'] for p in pairs) if pairs else 0

    wb = Workbook()

    # ── Tab 1: Summary ──
    ws1 = wb.active; ws1.title = "规律性分析结果"
    ws1.merge_cells('A1:H1'); ws1['A1'] = "技术标规律性分析检查结果"
    ws1['A1'].font = Font(bold=True, size=14)
    ws1['A2'] = "标段名称：用户自定义"
    ws1['A3'] = f"投标单位个数：{n}"
    ws1['A4'] = f"创建时间：{_ts}"
    ws1['A5'] = (
        f"检查结果：文本相似度检查{'存在异常' if max_sim > 0 else '无异常'}；"
        "重点信息无异常；文档属性检查无异常；图片相似度检查无异常"
    )
    ws1['A6'] = (
        "检查规则：检查相似度≥80%的段落，文本中重点信息，相似图片，相同作者；"
        "忽略与招标文件相同内容，忽略标点符号及小于6个字的内容，忽略目录，忽略文件中的技术标准"
    )
    ws1.merge_cells('A6:H6')

    row = 10
    ws1[f'A{row}'] = "一、标书围串风险分析结果"
    ws1[f'A{row}'].font = Font(bold=True)
    row += 1
    headers = ["投标单位"] + [fd['filename'] for fd in file_data]
    for col, h in enumerate(headers, 1):
        c = ws1.cell(row=row, column=col, value=h)
        c.font = Font(bold=True); c.alignment = Alignment(horizontal='center')
    row += 1
    for i in range(n):
        ws1.cell(row=row, column=1, value=file_data[i]['filename'])
        for j in range(n):
            val = "--" if i == j else next(
                (p['risk'] for p in pairs
                 if (p['i'] == i and p['j'] == j) or (p['i'] == j and p['j'] == i)),
                0)
            ws1.cell(row=row, column=j + 2, value=val)
        row += 1
    row += 2
    ws1[f'A{row}'] = "二、分析结果详情"
    ws1[f'A{row}'].font = Font(bold=True)
    row += 1
    detail_headers = ["序号", "投标单位1", "投标单位2", "风险度",
                      "文本相似度（%）", "语义相似度（%）", "图片相似度（%）",
                      "文件属性雷同", "重点信息雷同（项）"]
    for col, h in enumerate(detail_headers, 1):
        c = ws1.cell(row=row, column=col, value=h)
        c.font = Font(bold=True); c.alignment = Alignment(horizontal='center')
    row += 1
    for idx, p in enumerate(pairs, 1):
        ws1.cell(row=row, column=1, value=idx)
        ws1.cell(row=row, column=2, value=p['name1'])
        ws1.cell(row=row, column=3, value=p['name2'])
        ws1.cell(row=row, column=4, value=p['risk'])
        ws1.cell(row=row, column=5, value=p['sim'])
        ws1.cell(row=row, column=6, value=p.get('semantic_sim', 0))
        ws1.cell(row=row, column=7, value=0)
        ws1.cell(row=row, column=8, value="是" if p['attr_same'] else "否")
        ki_match = next(
            (k for k in key_info_matches
             if k['name1'] == p['name1'] and k['name2'] == p['name2']),
            None)
        ws1.cell(row=row, column=9, value=len(ki_match['common_keywords']) if ki_match else 0)
        row += 1
    for col in range(1, 10):
        ws1.column_dimensions[get_column_letter(col)].width = 20

    # ── Tab 2: Text detail matches ──
    ws2 = wb.create_sheet("规律性分析详情（文本）")
    ws2['A1'] = "规律性分析详情（文本）"
    ws2.merge_cells('A1:I1')
    ws2['A5'] = "序号"; ws2['B5'] = "目标单位"; ws2['C5'] = "目标单位对应文档"
    ws2['D5'] = "页码"; ws2['E5'] = "目标单位内容"
    ws2['F5'] = "对比单位"; ws2['G5'] = "对比单位对应文档"
    ws2['H5'] = "页码"; ws2['I5'] = "对比单位相似内容"
    for col in range(1, 10):
        ws2.cell(row=5, column=col).font = Font(bold=True)
    r2 = 6; seq = 1
    for p in pairs:
        if not p['blocks']:
            continue
        for blk in p['blocks']:
            ws2.cell(row=r2, column=1, value=seq)
            ws2.cell(row=r2, column=2, value=p['name1'])
            ws2.cell(row=r2, column=3, value=p['name1'] + ".pdf")
            ws2.cell(row=r2, column=4, value=blk.get('page1', ''))
            ws2.cell(row=r2, column=5, value=blk['text1_snippet'])
            ws2.cell(row=r2, column=6, value=p['name2'])
            ws2.cell(row=r2, column=7, value=p['name2'] + ".pdf")
            ws2.cell(row=r2, column=8, value=blk.get('page2', ''))
            ws2.cell(row=r2, column=9, value=blk['text2_snippet'])
            r2 += 1; seq += 1
    for col in range(1, 10):
        ws2.column_dimensions[get_column_letter(col)].width = 30

    # ── Tab 3: Key info matches ──
    ws3 = wb.create_sheet("规律性分析详情（重点信息）")
    ws3['A1'] = "规律性分析详情（重点信息）"
    ws3.merge_cells('A1:I1')
    ws3['A5'] = "序号"; ws3['B5'] = "AI识别类型"; ws3['C5'] = "内容"
    ws3['D5'] = "目标单位"; ws3['E5'] = "目标单位对应文档"
    ws3['F5'] = "页码"; ws3['G5'] = "对比单位"
    ws3['H5'] = "对比单位对应文档"; ws3['I5'] = "页码"
    for col in range(1, 10):
        ws3.cell(row=5, column=col).font = Font(bold=True)
    r3 = 6; seq3 = 1
    for ki in key_info_matches:
        for kw in ki['common_keywords']:
            ws3.cell(row=r3, column=1, value=seq3)
            ws3.cell(row=r3, column=2, value="关键词")
            ws3.cell(row=r3, column=3, value=kw)
            ws3.cell(row=r3, column=4, value=ki['name1'])
            ws3.cell(row=r3, column=5, value=ki['name1'] + ".pdf")
            ws3.cell(row=r3, column=7, value=ki['name2'])
            ws3.cell(row=r3, column=8, value=ki['name2'] + ".pdf")
            r3 += 1; seq3 += 1
    for col in range(1, 10):
        ws3.column_dimensions[get_column_letter(col)].width = 20

    return wb


# ═══════════════════════════════════════════════════════════════════════════
# HTML report builder
# ═══════════════════════════════════════════════════════════════════════════

def build_summary_html(file_data, pairs, risk_matrix):
    """Build the collapsible summary section HTML."""
    n = len(file_data)
    high_risk_files = []
    strong_alert_files = []
    for i in range(n):
        for j in range(i + 1, n):
            if risk_matrix[i][j] > 20:
                strong_alert_files.extend([file_data[i]['filename'], file_data[j]['filename']])
            elif risk_matrix[i][j] > 10:
                high_risk_files.extend([file_data[i]['filename'], file_data[j]['filename']])
    strong_alert_files = list(set(strong_alert_files))
    high_risk_files = list(set(high_risk_files) - set(strong_alert_files))

    html_out = (
        '<details style="margin-bottom:4px; border-radius:6px; padding:6px;">'
        '<summary style="cursor:pointer; font-weight:bold; font-size:0.9rem;">'
        '📋 对比摘要 (点击展开)</summary>'
        '<div style="margin-top:12px; border-left:8px solid #2c3e50; padding-left:8px;">'
    )
    for fd in file_data:
        preview = html.escape(fd['text'][:200].replace('\n', ' ')) + '…'
        safe_fn = html.escape(fd['filename'])
        html_out += (
            f'<div style="margin-bottom:15px;">'
            f'<strong>📄 {safe_fn}</strong><br>'
            f'<span style="color:#666; font-size:0.85rem;">{preview}</span></div>'
        )
    if strong_alert_files:
        safe_strong = ', '.join(html.escape(f) for f in strong_alert_files)
        html_out += (
            f'<p style="color:#d9534f; font-weight:bold;">'
            f'🚨 强烈警告：以下文件风险度超过20：{safe_strong}</p>'
        )
    elif high_risk_files:
        safe_high = ', '.join(html.escape(f) for f in high_risk_files)
        html_out += (
            f'<p style="color:#f0ad4e; font-weight:bold;">'
            f'⚠️ 可疑文件：以下文件风险度超过10：{safe_high}</p>'
        )
    else:
        html_out += '<p style="color:#5cb85c;">✅ 未发现高风险文件（风险度均≤10）</p>'
    html_out += '</div></details>'
    return html_out


def build_pair_report_html(file_data, pairs, risk_matrix):
    """Build the detailed pair report or risk matrix HTML."""
    n = len(file_data)
    short_names = [truncate_filename(fd['filename'], 20) for fd in file_data]

    if n == 2:
        p = pairs[0]
        if p['blocks']:
            detail_rows = ""
            for b in p['blocks']:
                detail_rows += (
                    f'<tr>'
                    f'<td style="border:1px solid #ccc; padding:8px; text-align:center;">{b["id"]}</td>'
                    f'<td style="border:1px solid #ccc; padding:8px; text-align:center;">{b["size"]}</td>'
                    f'<td style="border:1px solid #ccc; padding:8px; word-break:break-word; max-width:300px;">{html.escape(b["text1_snippet"])}</td>'
                    f'<td style="border:1px solid #ccc; padding:8px; word-break:break-word; max-width:300px;">{html.escape(b["text2_snippet"])}</td>'
                    f'</tr>'
                )
            return (
                f'<details><summary style="cursor:pointer; font-weight:bold;">'
                f'📋 详细相似度明细报告（共 {len(p["blocks"])} 个匹配块）</summary>'
                f'<div style="margin-top:12px;">'
                f'<p><strong>总匹配字符数：</strong>{sum(b["size"] for b in p["blocks"])} 字符 &nbsp;|&nbsp;'
                f'<strong>平均匹配块长度：</strong>{round(sum(b["size"] for b in p["blocks"]) / len(p["blocks"]), 1)} 字符</p>'
                f'<div style="overflow-x:auto;">'
                f'<table style="width:100%; border-collapse:collapse; margin-top:10px;">'
                f'<thead><tr style="background:#f0f0f0;">'
                f'<th style="border:1px solid #ccc; padding:8px;">块序号</th>'
                f'<th style="border:1px solid #ccc; padding:8px;">匹配字符数</th>'
                f'<th style="border:1px solid #ccc; padding:8px;">文档A片段</th>'
                f'<th style="border:1px solid #ccc; padding:8px;">文档B片段</th>'
                f'</tr></thead><tbody>{detail_rows}</tbody></table></div></div></details>'
            )
        return "<p>未检测到显著匹配块。</p>"
    else:
        matrix_html = (
            '<details><summary style="cursor:pointer; font-weight:bold;">'
            '📊 风险度矩阵 (点击展开/折叠)</summary>'
            '<div style="overflow-x:auto; margin-top:12px;">'
            '<table style="border-collapse:collapse; font-size:0.85rem; min-width:400px; width:100%;">'
            f'<thead><tr><th style="padding:8px; border:1px solid #ddd;"></th>'
            + ''.join(f'<th style="padding:8px; border:1px solid #ddd; word-break:break-word;">{html.escape(short_names[i])}</th>' for i in range(n))
            + '</tr></thead><tbody>'
        )
        for i in range(n):
            matrix_html += (
                f'<tr><td style="border:1px solid #ddd; padding:8px; font-weight:bold;">'
                f'{html.escape(short_names[i])}</td>'
            )
            for j in range(n):
                if i == j:
                    val, bg = '--', ''
                else:
                    val = f'{risk_matrix[i][j]:.2f}'
                    if risk_matrix[i][j] > 20:
                        bg = ' style="background:#d9534f; color:white; font-weight:bold;"'
                    elif risk_matrix[i][j] > 10:
                        bg = ' style="background:#f0ad4e;"'
                    else:
                        bg = ''
                matrix_html += (
                    f'<td style="border:1px solid #ddd; padding:8px; text-align:center;"{bg}>'
                    f'{html.escape(val)}</td>'
                )
            matrix_html += '</tr>'
        matrix_html += (
            '</tbody></table></div>'
            '<p style="font-size:0.7rem; color:#666; margin-top:8px;">风险度矩阵（值越高风险越大）</p>'
            '</details>'
        )
        return matrix_html


def build_full_report_html(file_data, pairs, risk_matrix,
                           typo_results=None, rel_report=None, quote_result=None,
                           ai_analysis_html=""):
    """Build the complete HTML report document.

    This is the full <!DOCTYPE html> document, NOT just a fragment —
    used for the ZIP download and the chat message.
    """
    n = len(file_data)
    _ts = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')
    max_risk = max(p['risk'] for p in pairs) if pairs else 0

    html_out = f"""<!DOCTYPE html><html lang="zh"><head><meta charset="UTF-8"><title>批量对比报告</title>
<style>body{{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;max-width:900px;margin:0 auto;padding:20px;color:#1e293b;line-height:1.6}}
h1{{color:#0f172a;border-bottom:3px solid #2563eb;padding-bottom:8px}}h2{{color:#334155;margin-top:24px}}
.card{{background:#f8fafc;border:1px solid #e2e8f0;border-radius:8px;padding:16px;margin:12px 0}}
.risk-high{{color:#dc2626;font-weight:bold}}.risk-warn{{color:#d97706;font-weight:bold}}.risk-ok{{color:#16a34a}}
table{{width:100%;border-collapse:collapse;margin:12px 0;font-size:.85rem}}th,td{{border:1px solid #e2e8f0;padding:8px;text-align:left}}
th{{background:#f1f5f9}}.highlight{{background:#fef9c3}}</style></head><body>
<h1>📁 批量对比报告</h1>
<div class="card"><strong>对比文件数：</strong>{n} &nbsp;|&nbsp; <strong>对比对数：</strong>{len(pairs)} &nbsp;|&nbsp; <strong>最高风险度：</strong>{max_risk:.1f} &nbsp;|&nbsp; <strong>生成时间：</strong>{_ts}</div>
<h2>📋 文件列表</h2><ul>{''.join(f'<li><strong>{html.escape(fd["filename"])}</strong></li>' for fd in file_data)}</ul>
<h2>📊 风险度矩阵</h2><table><tr><th></th>{''.join(f'<th>{html.escape(truncate_filename(fd["filename"], 15))}</th>' for fd in file_data)}</tr>
{''.join(f'<tr><td><strong>{html.escape(truncate_filename(file_data[i]["filename"], 15))}</strong></td>{"".join("<td style=\"text-align:center;\">--</td>" if i==j else f"<td class=\"{'risk-high' if risk_matrix[i][j]>20 else 'risk-warn' if risk_matrix[i][j]>10 else 'risk-ok'}\">{risk_matrix[i][j]:.1f}</td>" for j in range(n))}</tr>' for i in range(n))}
</table>"""

    # Warnings
    high_risk_files = []
    strong_alert_files = []
    for i in range(n):
        for j in range(i + 1, n):
            if risk_matrix[i][j] > 20:
                strong_alert_files.extend([file_data[i]['filename'], file_data[j]['filename']])
            elif risk_matrix[i][j] > 10:
                high_risk_files.extend([file_data[i]['filename'], file_data[j]['filename']])
    strong_alert_files = list(set(strong_alert_files))
    high_risk_files = list(set(high_risk_files) - set(strong_alert_files))
    short_names = [truncate_filename(fd['filename'], 15) for fd in file_data]

    if strong_alert_files:
        html_out += (
            f'<div class="card" style="background:#fef2f2;border-color:#fecaca;">'
            f'<h3 style="margin-top:0;color:#dc2626;">🚨 强烈警告</h3>'
            f'<p>以下文件风险度超过20：<strong>{", ".join(html.escape(f) for f in strong_alert_files)}</strong></p></div>'
        )
    elif high_risk_files:
        html_out += (
            f'<div class="card" style="background:#fffbeb;border-color:#fde68a;">'
            f'<h3 style="margin-top:0;color:#d97706;">⚠️ 需关注</h3>'
            f'<p>以下文件风险度超过10：<strong>{", ".join(html.escape(f) for f in high_risk_files)}</strong></p></div>'
        )

    # Pair cards
    for p in pairs[:10]:
        risk_class = 'risk-high' if p['risk'] > 20 else ('risk-warn' if p['risk'] > 10 else 'risk-ok')
        html_out += (
            f'<div class="card"><h3>📄 {html.escape(p["name1"])} ↔ {html.escape(p["name2"])}</h3>'
            f'<p>风险度: <span class="{risk_class}">{p["risk"]:.1f}</span> | '
            f'文本相似度: {p["sim"]:.1f}% | 相同作者: {"是" if p["attr_same"] else "否"}</p>'
        )
        if p['blocks']:
            html_out += f'<p>匹配段落数: {len(p["blocks"])} · 总匹配字符: {sum(b["size"] for b in p["blocks"])}</p>'
        html_out += '</div>'

    # Typo section
    if typo_results:
        total_typos = sum(r.total_suspects for r in typo_results.values())
        total_crit = sum(r.critical_count for r in typo_results.values())
        if total_typos > 0:
            html_out += '<h2>📝 错别字检测</h2>'
            html_out += (
                f'<div class="card"><p><strong>疑似错别字:</strong> {total_typos} 处 | '
                f'<strong>严重:</strong> {total_crit} 处</p>'
                '<table style="font-size:0.85rem;">'
                '<tr><th>文件</th><th>层次</th><th>疑似文本</th><th>建议</th><th>置信度</th><th>严重性</th></tr>'
            )
            for doc_name, report in typo_results.items():
                for f in report.findings[:30]:
                    sev_class = 'risk-high' if f.severity == 'critical' else ('risk-warn' if f.severity == 'warning' else '')
                    html_out += (
                        f'<tr><td>{html.escape(doc_name[:20])}</td>'
                        f'<td>{f.layer}</td>'
                        f'<td><code>{html.escape(f.suspect_text[:40])}</code></td>'
                        f'<td>{html.escape(", ".join(f.suggestions[:3]) if f.suggestions else "—")}</td>'
                        f'<td>{f.confidence:.0%}</td>'
                        f'<td class="{sev_class}">{f.severity}</td></tr>'
                    )
            html_out += '</table>'
            if total_typos > 30:
                html_out += (
                    f'<p style="color:#64748b;font-size:.85rem;">'
                    f'（仅显示前30项，共{total_typos}项）</p>'
                )
            html_out += '</div>'

    # Relationship section
    if rel_report and rel_report.red_flags:
        html_out += '<h2>🔗 关联关系分析</h2>'
        html_out += (
            f'<div class="card"><p><strong>提取实体:</strong> {len(rel_report.entities)} | '
            f'<strong>发现关系:</strong> {len(rel_report.relationships)} | '
            f'<strong>风险评分:</strong> <span class="{("risk-high" if rel_report.risk_score > 50 else ("risk-warn" if rel_report.risk_score > 20 else "risk-ok"))}">{rel_report.risk_score:.1f}</span></p>'
            f'<p><strong>检测模块:</strong> {", ".join(rel_report.modules_run)}</p>'
        )
        if rel_report.red_flags:
            html_out += '<ul>'
            for flag in rel_report.red_flags[:15]:
                html_out += f'<li class="risk-warn">{html.escape(flag)}</li>'
            html_out += '</ul>'
        cpm = rel_report.company_personnel_map
        if cpm.get('manual_review_required') and cpm.get('companies'):
            html_out += (
                '<details><summary style="cursor:pointer;font-weight:bold;margin-top:8px;">'
                '📋 公司与关键人员清单（供人工审查）</summary>'
                '<table style="margin-top:8px;"><tr><th>公司名称</th><th>关键人员</th><th>涉及文件</th></tr>'
            )
            for comp in cpm['companies'][:20]:
                personnel_str = '; '.join(f"{p['name']}({p['title']})" for p in comp['personnel'][:5])
                html_out += (
                    f'<tr><td>{html.escape(comp["name"])}</td>'
                    f'<td>{html.escape(personnel_str)}</td>'
                    f'<td>{comp["file_count"]}个文件</td></tr>'
                )
            html_out += '</table></details>'
        html_out += '</div>'

    # Quote section
    if quote_result and quote_result.get('per_bidder'):
        html_out += '<h2>💰 报价异常检测</h2>'
        html_out += (
            '<div class="card"><table>'
            '<tr><th>投标单位</th><th>风险评分</th><th>离散系数(CV)</th>'
            '<th>同价疑义</th><th>异常降幅</th><th>聚类疑义</th><th>本福特偏差</th></tr>'
        )
        for pb in quote_result['per_bidder']:
            flags = []
            if pb.get('same_rate_flag'): flags.append('⚠️同价')
            if pb.get('abnormal_drop_flag'): flags.append('⬇️异常降幅')
            if pb.get('clustering_flag'): flags.append('🔗聚类')
            flag_str = ', '.join(flags) if flags else '✅ 正常'
            risk_class = 'risk-high' if pb.get('risk_score', 0) > 50 else ('risk-warn' if pb.get('risk_score', 0) > 20 else 'risk-ok')
            html_out += (
                f'<tr><td>{html.escape(pb["filename"])}</td>'
                f'<td class="{risk_class}">{pb.get("risk_score", 0):.1f}</td>'
                f'<td>{pb.get("cv", 0):.4f}</td>'
                f'<td>{"是" if pb.get("same_rate_flag") else "否"}</td>'
                f'<td>{"是" if pb.get("abnormal_drop_flag") else "否"}</td>'
                f'<td>{"是" if pb.get("clustering_flag") else "否"}</td>'
                f'<td>{pb.get("benford_deviation", 0):.3f}</td></tr>'
            )
        html_out += '</table>'
        if quote_result.get('cross_same_rate'):
            html_out += '<p class="risk-warn">⚠️ 跨投标单位同价疑义：多个投标单位首轮报价异常接近</p>'
        if quote_result.get('cross_clustering'):
            html_out += '<p class="risk-warn">🔗 跨投标单位价格聚类：多个投标单位报价集中在异常窄区间</p>'
        html_out += (
            f'<p style="color:#64748b;font-size:.85rem;">'
            f'最高报价风险评分: {quote_result.get("max_risk_score", 0):.1f} | '
            f'平均CV: {quote_result.get("avg_cv", 0):.4f}</p>'
        )
        html_out += '</div>'

    html_out += '<p style="margin-top:24px;color:#64748b;font-size:.85rem;">完整风险矩阵、文本匹配详情、重点信息雷同、文件属性分析请参见配套DOCX文件。</p>'
    if ai_analysis_html:
        html_out += ai_analysis_html
    html_out += '</body></html>'
    return html_out


# ═══════════════════════════════════════════════════════════════════════════
# DOCX report builder (replaces Excel)
# ═══════════════════════════════════════════════════════════════════════════

def build_report_docx(file_data, pairs, key_info_matches, attr_details=None):
    """Build a .docx comparison report following professional bid-review formatting.

    Returns bytes ready to write to a .docx file.
    """
    n = len(file_data)
    max_risk = max(p['risk'] for p in pairs) if pairs else 0
    _ts = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')
    risk_matrix = [[0] * n for _ in range(n)]
    for p in pairs:
        risk_matrix[p['i']][p['j']] = p['risk']
        risk_matrix[p['j']][p['i']] = p['risk']

    if attr_details is None:
        attr_details = build_attr_details(file_data)

    def _safe(s):
        return str(s) if s else ''

    doc = DocxDocument()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10)

    thin_border = {'val': 'single', 'sz': '4', 'color': 'CCCCCC'}

    # ── Cover page ──
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('批量对比分析报告')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Bid Document Comparison Report')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()
    meta_items = [
        ('分析日期', _ts),
        ('对比文件数', str(n)),
        ('对比对数', str(len(pairs))),
        ('最高风险度', f'{max_risk:.1f}'),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}: {value}')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x47, 0x56, 0x69)

    doc.add_page_break()

    # ── Section 1: Risk matrix ──
    h = doc.add_paragraph()
    run = h.add_run('一、风险矩阵总览')
    run.bold = True
    run.font.size = Pt(14)

    table = doc.add_table(rows=n + 1, cols=n + 1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.cell(0, 0).text = ''
    for j, fd in enumerate(file_data):
        table.cell(0, j + 1).text = truncate_filename(fd['filename'], 15)
    for i in range(n):
        table.cell(i + 1, 0).text = truncate_filename(file_data[i]['filename'], 15)
        for j in range(n):
            cell = table.cell(i + 1, j + 1)
            if i == j:
                cell.text = '--'
            else:
                val = risk_matrix[i][j]
                cell.text = f'{val:.1f}'
                if val > 20:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                            run.bold = True

    # Warnings
    strong_alert_files = []
    high_risk_files = []
    for i in range(n):
        for j in range(i + 1, n):
            if risk_matrix[i][j] > 20:
                strong_alert_files.extend([file_data[i]['filename'], file_data[j]['filename']])
            elif risk_matrix[i][j] > 10:
                high_risk_files.extend([file_data[i]['filename'], file_data[j]['filename']])
    strong_alert_files = list(set(strong_alert_files))
    high_risk_files = list(set(high_risk_files) - set(strong_alert_files))

    doc.add_paragraph()
    if strong_alert_files:
        p = doc.add_paragraph()
        run = p.add_run(f'🚨 强烈警告: {", ".join(strong_alert_files)} 风险度超过20')
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        run.bold = True
    elif high_risk_files:
        p = doc.add_paragraph()
        run = p.add_run(f'⚠️ 需关注: {", ".join(high_risk_files)} 风险度超过10')
        run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
        run.bold = True

    doc.add_page_break()

    # ── Section 2: Detail results ──
    h = doc.add_paragraph()
    run = h.add_run('二、对比结果详情')
    run.bold = True
    run.font.size = Pt(14)

    detail_table = doc.add_table(rows=len(pairs) + 1, cols=5)
    detail_table.style = 'Table Grid'
    detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['序号', '单位A', '单位B', '风险度', '相似度(%)']
    for j, hdr in enumerate(headers):
        cell = detail_table.cell(0, j)
        cell.text = hdr
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for idx, p in enumerate(pairs):
        detail_table.cell(idx + 1, 0).text = str(idx + 1)
        detail_table.cell(idx + 1, 1).text = truncate_filename(p['name1'], 20)
        detail_table.cell(idx + 1, 2).text = truncate_filename(p['name2'], 20)
        detail_table.cell(idx + 1, 3).text = f'{p["risk"]:.1f}'
        detail_table.cell(idx + 1, 4).text = f'{p["sim"]:.1f}'

    doc.add_page_break()

    # ── Section 3: Text matching blocks ──
    has_blocks = any(p['blocks'] for p in pairs)
    if has_blocks:
        h = doc.add_paragraph()
        run = h.add_run('三、文本雷同详情')
        run.bold = True
        run.font.size = Pt(14)

        for pair in pairs:
            if not pair['blocks']:
                continue
            doc.add_paragraph()
            sub = doc.add_paragraph()
            run = sub.add_run(f'{truncate_filename(pair["name1"], 20)} ↔ {truncate_filename(pair["name2"], 20)}')
            run.font.size = Pt(11)

            block_table = doc.add_table(rows=min(len(pair['blocks']), 50) + 1, cols=4)
            block_table.style = 'Table Grid'
            block_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, hdr in enumerate(['序号', '页码', '文档A片段', '文档B片段']):
                cell = block_table.cell(0, j)
                cell.text = hdr
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(8)

            for bi, blk in enumerate(pair['blocks'][:50]):
                block_table.cell(bi + 1, 0).text = str(blk.get('id', bi + 1))
                block_table.cell(bi + 1, 1).text = _safe(blk.get('page1', ''))
                block_table.cell(bi + 1, 2).text = _safe(blk.get('text1_snippet', ''))[:200]
                block_table.cell(bi + 1, 3).text = _safe(blk.get('text2_snippet', ''))[:200]

            if len(pair['blocks']) > 50:
                p = doc.add_paragraph()
                run = p.add_run(f'（仅显示前50项，共{len(pair["blocks"])}项）')
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_page_break()

    # ── Section 4: Key info matches ──
    h = doc.add_paragraph()
    run = h.add_run('四、重点信息雷同')
    run.bold = True
    run.font.size = Pt(14)

    if key_info_matches:
        ki_table = doc.add_table(rows=len(key_info_matches) + 1, cols=3)
        ki_table.style = 'Table Grid'
        ki_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, hdr in enumerate(['序号', '单位A', '共同关键词']):
            cell = ki_table.cell(0, j)
            cell.text = hdr
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for ki_idx, ki in enumerate(key_info_matches):
            ki_table.cell(ki_idx + 1, 0).text = str(ki_idx + 1)
            ki_table.cell(ki_idx + 1, 1).text = f'{ki["name1"]} ↔ {ki["name2"]}'
            ki_table.cell(ki_idx + 1, 2).text = ', '.join(ki.get('common_keywords', [])[:15])
    else:
        p = doc.add_paragraph()
        run = p.add_run('未发现重点信息雷同。')
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_page_break()

    # ── Section 5: File attributes ──
    h = doc.add_paragraph()
    run = h.add_run('五、文件属性对比')
    run.bold = True
    run.font.size = Pt(14)

    attr_table = doc.add_table(rows=len(attr_details) + 1, cols=4)
    attr_table.style = 'Table Grid'
    attr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, hdr in enumerate(['文件名', '作者', '创建日期', '创建程序']):
        cell = attr_table.cell(0, j)
        cell.text = hdr
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for ai, ad in enumerate(attr_details):
        attr_table.cell(ai + 1, 0).text = truncate_filename(ad.get('filename', ''), 25)
        attr_table.cell(ai + 1, 1).text = _safe(ad.get('author', ''))
        attr_table.cell(ai + 1, 2).text = _safe(ad.get('creation_date', ''))
        attr_table.cell(ai + 1, 3).text = _safe(ad.get('creator', ''))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
    return html_out
