"""Per-user writing style profiling engine.

Analyzes chat history to extract writing style metrics:
- Average sentence length (chars)
- Formality score (0–100)
- Top frequent keywords (Chinese tokenized)
- Preferred tone indicators

Stores persistent profiles at data/user_styles/{user_id}.json.
Profiles evolve iteratively — each analysis merges with prior data.
"""

import os, re, json, logging
from collections import Counter
from datetime import datetime, timezone
from threading import Lock
from pathlib import Path

logger = logging.getLogger(__name__)

DATA_DIR = Path(__file__).parent.parent.parent / "data"
STYLES_DIR = DATA_DIR / "user_styles"
_lock = Lock()

# ── Analysis configuration ──
FORMAL_WORDS = {
    '请', '您', '贵', '谨', '特此', '兹', '据此', '函告', '批复', '请示',
    '审批', '备案', '招标', '投标', '合同', '协议', '条款', '履行',
    '确认', '核实', '审核', '评估', '报告', '通知', '上述', '如下',
    '予以', '予以', '相关', '标准化', '规范化', '合规', '资质',
    '预算', '结算', '验收', '交付', '期限', '截止',
}
INFORMAL_WORDS = {
    '吗', '吧', '呢', '呀', '哈', '嗯', '哦', '啥', '咋', '咋办',
    '行不', '行吗', '可以吗', '能不能', '帮我看下', '怎么看',
    '大概', '差不多', '好像', '应该', '感觉', '觉得', '可能',
}
SENTENCE_SPLIT = re.compile(r'[。！？；\n]')

# POS tag categories for tone/domain analysis
NOUN_TAGS = {'n', 'nr', 'ns', 'nt', 'nz', 'ng'}       # nouns, proper names, locations
VERB_TAGS = {'v', 'vd', 'vn', 'vx'}                    # verbs
ADJ_TAGS  = {'a', 'ad', 'an'}                          # adjectives
FUNC_TAGS = {'p', 'c', 'u', 'e', 'y', 'o', 'w', 'x'}  # particles, conjunctions, etc.
DOMAIN_TAGS = {'nz', 'eng'}                             # domain-specific terms, English


def _ensure_dirs():
    os.makedirs(STYLES_DIR, exist_ok=True)


# ── Analysis functions ──

def _analyze_sentences(messages: list[str]) -> dict:
    """Analyze sentence structure from a list of user messages."""
    all_text = ' '.join(messages)
    sentences = [s.strip() for s in SENTENCE_SPLIT.split(all_text) if s.strip() and len(s.strip()) > 2]
    if not sentences:
        return {'avg_sentence_chars': 0, 'total_sentences': 0}

    lengths = [len(s) for s in sentences]
    avg_len = round(sum(lengths) / len(lengths), 1)

    # Classify: short (<30), medium (30–80), long (>80)
    short = sum(1 for l in lengths if l < 30)
    medium = sum(1 for l in lengths if 30 <= l <= 80)
    long = sum(1 for l in lengths if l > 80)

    return {
        'avg_sentence_chars': avg_len,
        'total_sentences': len(sentences),
        'short_pct': round(short / len(sentences) * 100, 1),
        'medium_pct': round(medium / len(sentences) * 100, 1),
        'long_pct': round(long / len(sentences) * 100, 1),
    }


def _analyze_formality(messages: list[str]) -> dict:
    """Score formality based on word choice ratio using jieba tokenization."""
    all_text = ''.join(messages)
    from app.services.text_utils import lcut as tokenize
    words = set(tokenize(all_text))  # unique words for boundary-accurate matching
    formal_count = sum(1 for w in FORMAL_WORDS if w in words)
    informal_count = sum(1 for w in INFORMAL_WORDS if w in words)
    total = formal_count + informal_count
    if total == 0:
        return {'formality_score': 50, 'formal_words_hit': 0, 'informal_words_hit': 0}

    score = round(formal_count / total * 100)
    return {
        'formality_score': score,
        'formal_words_hit': formal_count,
        'informal_words_hit': informal_count,
    }


def _analyze_pos_distribution(messages: list[str]) -> dict:
    """Analyze POS tag distribution using jieba.posseg for deeper style insight."""
    all_text = ''.join(messages)
    tagged = posseg(all_text)
    if not tagged:
        return {}

    tags = Counter(flag for _, flag in tagged)
    noun_pct = round(sum(tags.get(t, 0) for t in NOUN_TAGS) / max(len(tagged), 1) * 100, 1)
    verb_pct = round(sum(tags.get(t, 0) for t in VERB_TAGS) / max(len(tagged), 1) * 100, 1)
    adj_pct  = round(sum(tags.get(t, 0) for t in ADJ_TAGS) / max(len(tagged), 1) * 100, 1)

    return {
        'total_pos_tags': len(tagged),
        'noun_pct': noun_pct,
        'verb_pct': verb_pct,
        'adj_pct': adj_pct,
        'domain_heavy': noun_pct > 40,  # noun-heavy = domain/technical user
    }


def _extract_keywords(messages: list[str], top_k: int = 20) -> list[dict]:
    """Extract top Chinese keywords from user messages using jieba segmentation."""
    all_text = ''.join(messages)
    from app.services.text_utils import top_keywords as tk_extract
    kw_list = tk_extract(all_text, top_k=top_k, min_len=2)
    return [{'word': w, 'count': c} for w, c in kw_list]


def _analyze_tone(messages: list[str]) -> dict:
    """Infer preferred tone from punctuation and sentence starters."""
    all_text = ''.join(messages)
    # Question ratio
    questions = sum(1 for m in messages if m.strip().endswith('？') or m.strip().endswith('?'))
    exclamations = sum(1 for m in messages if '！' in m or '!' in m)
    # Bullet-point style
    bullet_lines = sum(1 for m in messages if re.match(r'^[\d一二三四五六七八九十]+[、．.]', m.strip()))

    return {
        'question_ratio': round(questions / max(len(messages), 1) * 100, 1),
        'bullet_style_ratio': round(bullet_lines / max(len(messages), 1) * 100, 1),
        'uses_bullet_points': bullet_lines > len(messages) * 0.1,
    }


# ── Profile management ──

def _get_profile_path(user_id: str) -> str:
    return os.path.join(STYLES_DIR, f'{user_id}.json')


def analyze_user_style(user_id: str, sample_size: int = 100) -> dict:
    """Analyze a user's chat writing style from DB history.

    Args:
        user_id: user to analyze
        sample_size: max number of recent messages to sample

    Returns the complete profile dict.
    """
    try:
        from app.database import get_db_connection
    except ImportError:
        logger.error("Cannot import database")
        return _empty_profile(user_id)

    try:
        with get_db_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT cm.content FROM chat_messages cm
                    JOIN chat_sessions cs ON cm.thread_id = cs.thread_id
                    WHERE cs.user_id = %s AND cm.role = 'user' AND cm.content IS NOT NULL
                    ORDER BY cm.timestamp DESC LIMIT %s
                """, (user_id, sample_size))
                rows = cur.fetchall()
    except Exception as e:
        logger.warning(f"DB query for style analysis failed (user {user_id}): {e}")
        return _empty_profile(user_id)

    messages = [r[0] for r in rows if r[0] and r[0].strip()]
    if not messages:
        return _empty_profile(user_id)

    prev = _load_profile(user_id)
    sentence = _analyze_sentences(messages)
    formality = _analyze_formality(messages)
    keywords = _extract_keywords(messages)
    tone = _analyze_tone(messages)
    pos_dist = _analyze_pos_distribution(messages)

    # Merge with prior data (weighted average: 70% new, 30% old)
    merge_sentence = sentence  # fresh analysis, no merge needed for counts
    merge_formality = formality
    if prev and prev.get('message_count', 0) > 0:
        old_weight = min(0.3, prev['message_count'] / (prev['message_count'] + len(messages)))
        new_weight = 1 - old_weight
        merge_sentence = {
            **sentence,
            'avg_sentence_chars': round(
                sentence['avg_sentence_chars'] * new_weight +
                prev.get('sentence', {}).get('avg_sentence_chars', sentence['avg_sentence_chars']) * old_weight, 1
            ),
        }
        merge_formality = {
            **formality,
            'formality_score': round(
                formality['formality_score'] * new_weight +
                prev.get('formality', {}).get('formality_score', formality['formality_score']) * old_weight
            ),
        }

    # Merge keywords (accumulate counts, keep top 30)
    kw_map = {k['word']: k['count'] for k in keywords}
    if prev and prev.get('keywords'):
        for pk in prev['keywords']:
            kw_map[pk['word']] = kw_map.get(pk['word'], 0) + pk['count']
    merged_kw = sorted(
        [{'word': w, 'count': c} for w, c in kw_map.items()],
        key=lambda x: x['count'], reverse=True
    )[:30]

    profile = {
        'user_id': user_id,
        'message_count': len(messages),
        'total_analyzed': (prev.get('total_analyzed', 0) + len(messages)) if prev else len(messages),
        'sentence': merge_sentence,
        'formality': merge_formality,
        'keywords': merged_kw,
        'tone': tone,
        'pos_distribution': pos_dist,
        'last_analyzed': datetime.now(timezone.utc).isoformat(),
        'version': (prev.get('version', 0) + 1) if prev else 1,
        'style_label': _label_style(merge_formality.get('formality_score', 50), merge_sentence.get('avg_sentence_chars', 50)),
        'style_description': _describe_style(merge_formality.get('formality_score', 50), merge_sentence.get('avg_sentence_chars', 50), tone),
    }

    _save_profile(user_id, profile)
    return profile


def _label_style(formality: int, avg_len: float) -> str:
    """Generate a concise style label."""
    f = 'Formal' if formality >= 60 else ('Semi-formal' if formality >= 35 else 'Casual')
    l = 'Concise' if avg_len < 30 else ('Detailed' if avg_len < 80 else 'Verbose')
    return f'{f} · {l}'


def _describe_style(formality: int, avg_len: float, tone: dict) -> str:
    """Generate a human-readable style description."""
    parts = []
    if formality >= 70:
        parts.append("Highly formal corporate language with precise terminology")
    elif formality >= 50:
        parts.append("Professional tone with balanced formality")
    else:
        parts.append("Conversational and approachable style")

    if avg_len < 30:
        parts.append("favors short, direct sentences")
    elif avg_len < 80:
        parts.append("uses medium-length structured sentences")
    else:
        parts.append("prefers detailed, comprehensive expressions")

    if tone.get('uses_bullet_points'):
        parts.append("frequently uses bullet points for clarity")
    return '. '.join(parts) + '.'


def _empty_profile(user_id: str) -> dict:
    return {
        'user_id': user_id,
        'message_count': 0,
        'total_analyzed': 0,
        'sentence': {'avg_sentence_chars': 0, 'total_sentences': 0, 'short_pct': 0, 'medium_pct': 0, 'long_pct': 0},
        'formality': {'formality_score': 50, 'formal_words_hit': 0, 'informal_words_hit': 0},
        'keywords': [],
        'tone': {'question_ratio': 0, 'bullet_style_ratio': 0, 'uses_bullet_points': False},
        'last_analyzed': None,
        'version': 0,
        'style_label': 'Not analyzed',
        'style_description': 'No chat data available for analysis.',
    }


# ── Profile I/O ──

def _load_profile(user_id: str) -> dict | None:
    """Load a user's style profile from disk."""
    path = _get_profile_path(user_id)
    if not os.path.exists(path):
        return None
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return None


def _save_profile(user_id: str, profile: dict):
    """Save or update a user's style profile to disk."""
    _ensure_dirs()
    path = _get_profile_path(user_id)
    with _lock:
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(profile, f, ensure_ascii=False, indent=2)


def get_user_style(user_id: str) -> dict:
    """Get a user's style profile (load from disk, or return empty)."""
    profile = _load_profile(user_id)
    return profile if profile else _empty_profile(user_id)


def update_user_style(user_id: str, updates: dict) -> dict:
    """Admin/user manually update style preferences (label, description overrides)."""
    profile = get_user_style(user_id)
    allowed = {'style_label', 'style_description'}
    for k, v in updates.items():
        if k in allowed and isinstance(v, str):
            profile[k] = v
    profile['last_edited'] = datetime.now(timezone.utc).isoformat()
    _save_profile(user_id, profile)
    return profile


def get_all_style_profiles() -> list[dict]:
    """Return all user style profiles (for admin)."""
    _ensure_dirs()
    profiles = []
    for f in os.listdir(STYLES_DIR):
        if f.endswith('.json'):
            user_id = f[:-5]
            profile = get_user_style(user_id)
            profiles.append(profile)
    profiles.sort(key=lambda p: p.get('total_analyzed', 0), reverse=True)
    return profiles


def delete_user_style(user_id: str) -> bool:
    """Delete a user's style profile (admin only)."""
    path = _get_profile_path(user_id)
    if os.path.exists(path):
        os.remove(path)
        return True
    return False


# ── Report format: Word (.docx) with .md fallback ──

def generate_report_file(report_text: str, filename_prefix: str, label: str,
                         style_profile: dict = None) -> str:
    """Generate a report file. Uses python-docx if available, falls back to .md.

    Args:
        report_text: the AI-generated report markdown text
        filename_prefix: e.g. 'weekly_report_2026-06-27'
        label: human-readable label for the report header
        style_profile: optional writing style profile for the user

    Returns the path to the generated file.
    """
    _ensure_dirs()
    try:
        import docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        return _generate_docx(report_text, filename_prefix, label, style_profile)
    except ImportError:
        logger.info("python-docx not installed — falling back to .md")
        return _generate_md(report_text, filename_prefix, label, style_profile)


def _generate_docx(text: str, prefix: str, label: str, style_profile: dict = None) -> str:
    """Generate a .docx report file."""
    import docx
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading(label, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')).font.size = Pt(9)

    # Style profile note
    if style_profile and style_profile.get('style_label'):
        style_para = doc.add_paragraph()
        style_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = style_para.add_run(f"Writing Style: {style_profile['style_label']}")
        run.font.size = Pt(8)
        run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)

    doc.add_paragraph()  # spacer

    # Parse markdown into paragraphs
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('|') and '|' in line[1:]:
            # Simple table cell
            doc.add_paragraph(line, style='No Spacing')
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)

    path = os.path.join(STYLES_DIR, f'{prefix}.docx')
    doc.save(path)
    return path


def _generate_md(text: str, prefix: str, label: str, style_profile: dict = None) -> str:
    """Generate a .md report file."""
    header = f"# {label}\n\nGenerated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}\n"
    if style_profile and style_profile.get('style_label'):
        header += f"\n*Writing Style: {style_profile['style_label']}*\n"
    header += "\n---\n\n"

    path = os.path.join(STYLES_DIR, f'{prefix}.md')
    with open(path, 'w', encoding='utf-8') as f:
        f.write(header + text)
    return path
