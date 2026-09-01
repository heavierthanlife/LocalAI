"""Document analysis service — runs all checkers and produces structured reports.

Provides:
  - run_analysis():      runs all available checkers, returns example-structured report
  - build_analysis_docx(): generates .docx following the bid-rigging clue analysis format
"""
import logging
from datetime import datetime, timezone
from io import BytesIO

from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from celery_app import celery as celery_app

from app.services.batch_orchestrator import (
    compute_all_pairs, build_key_info_matches, build_attr_details,
    RiskScorer, truncate_filename,
)
from app.services.indicator_defs import INDICATOR_DEFS, rule_text

logger = logging.getLogger(__name__)

# ── 权重复合指数 (FIX-010)：每指标权重 × 归一化得分 → 0-100 复合指数 ──
# 触发类权重最高，报价规律/专家一致性次之，基础/扩展类低。skip 指标权重 0。
INDICATOR_WEIGHTS = {
    # 触发指标
    'same_machine_code': 0.10, 'same_file_code': 0.10, 'same_dongle': 0.10,
    'tech_section_similar': 0.10, 'contact_person_same': 0.04,
    'economic_error_similar': 0.04, 'bid_ip_same': 0.10, 'decrypt_ip_same': 0.10,
    'download_ip_same': 0.10, 'cross_file_code_same': 0.06,
    'cross_contact_same': 0.03,
    # 核心指标
    'tender_query': 0.06, 'candidate_give_up': 0.08,
    'subjective_expert_spread': 0.08, 'low_win_rate': 0.05, 'high_win_rate': 0.05,
    'bidder_count_abnormal': 0.08, 'upload_interval_abnormal': 0.06,
    'subjective_expert_units': 0.08, 'specific_expert_score': 0.06,
    'clique_expert_scoring': 0.08, 'clique_expert_consistency': 0.08,
    'objective_score_abnormal': 0.05, 'high_price_abnormal': 0.08,
    'extension_abnormal': 0.05, 'tender_fail_abnormal': 0.04,
    'waste_rate_abnormal': 0.06, 'download_no_bid': 0.05, 'no_show_abnormal': 0.05,
    'expert_deviation_abnormal': 0.08,
    # 基础指标
    'cross_bid_ip': 0.06, 'cross_decrypt_ip': 0.06, 'cross_machine_code': 0.06,
    'bidder_agent_contact': 0.03, 'expert_tenderer_closeness': 0.04,
    'expert_agent_closeness': 0.04, 'cross_download_ip': 0.06,
    'cross_dongle': 0.06, 'tech_seal_check': 0.03,
    'expert_bidder_closeness': 0.03, 'bad_expert_score': 0.04,
    # 扩展指标
    'tech_score_abnormal': 0.05, 'commercial_score_abnormal': 0.05,
    'quote_proportional_float': 0.10, 'contact_phone_abnormal': 0.08,
}
DEFAULT_INDICATOR_WEIGHT = 0.03

# 每指标得分上限（归一化用）；超过视为已触发
INDICATOR_SCORE_CAPS = {
    'same_machine_code': 30, 'same_file_code': 30, 'same_dongle': 30,
    'tech_section_similar': 30, 'contact_person_same': 40, 'economic_error_similar': 15,
    'bid_ip_same': 30, 'decrypt_ip_same': 30, 'download_ip_same': 30,
    'cross_file_code_same': 30, 'cross_contact_same': 40,
    'candidate_give_up': 45, 'high_price_abnormal': 30, 'waste_rate_abnormal': 40,
    'bidder_count_abnormal': 30, 'expert_deviation_abnormal': 35,
    'subjective_expert_spread': 35, 'subjective_expert_units': 35,
    'clique_expert_scoring': 30, 'clique_expert_consistency': 30,
    'contact_phone_abnormal': 30, 'quote_proportional_float': 30,
    'tech_score_abnormal': 35, 'commercial_score_abnormal': 35,
    'extension_abnormal': 10,
}
DEFAULT_SCORE_CAP = 30.0


def _weighted_total_score(indicators: list[dict]) -> float:
    """0-100 权重复合指数：Σ(w_i × min(score_i/cap_i, 1)) / Σ(w_i) × 100.

    text_sim 三指标去重：same_file_code 全计，tech_section_similar /
    cross_file_code_same 按 0.3× 衰减（避免同一次相似度被计 3 次）。
    """
    # text_sim 指标去重：same_file_code 代表 text_sim 维度，
    # tech_section_similar / cross_file_code_same 不再单独贡献（避免同一次相似度计 3 次）
    text_sim_primary = 'same_file_code'
    num = 0.0
    den = 0.0
    for ind in indicators:
        if ind.get('skipped'):
            continue
        iid = ind.get('id', '')
        if iid in ('tech_section_similar', 'cross_file_code_same'):
            continue  # 去重：由 same_file_code 代表
        w = INDICATOR_WEIGHTS.get(iid, DEFAULT_INDICATOR_WEIGHT)
        cap = INDICATOR_SCORE_CAPS.get(iid, DEFAULT_SCORE_CAP)
        s = float(ind.get('score', 0) or 0)
        norm = min(s / cap, 1.0) if cap > 0 else 0.0
        num += w * norm
        den += w
    return round(num / den * 100, 1) if den > 0 else 0.0


def _run_checker(name, file_data, user_id, thread_id, tender_text=None):
    """Run a single checker with try-except, returns dict or None."""
    try:
        if name == 'skip':
            # Network/data-source dependent indicator — placeholder per design.
            return {'skipped': True, 'error': '需外部数据源（交易平台/评标系统数据）'}

        if name == 'text_sim':
            # 未提供招标文件时，无法去除招标模板 → 高余弦多是模板重叠，不是围标证据。
            # 标记为占位（诚实不误报），避免正常投标被标为"文本相似异常"。
            if not tender_text:
                return {'skipped': True, 'error': '需招标文件做模板去除，文本相似度不作依据'}
            from app.services.batch_compare_svc import _precompute_tfidf_for_files
            tfidf_matrix = None
            try:
                _vec, tfidf_matrix = _precompute_tfidf_for_files(file_data, template_text=tender_text)
            except Exception as e:
                logger.warning(f"text_sim tfidf precompute failed: {e}")
            pairs, risk_matrix = compute_all_pairs(file_data, {
                'text_sim': True, 'key_info': False,
                'file_attr': False, 'image_sim': False
            }, tfidf_matrix=tfidf_matrix, template_text=tender_text)
            max_risk = max(p['risk'] for p in pairs) if pairs else 0
            max_sim = max(p['sim'] for p in pairs) if pairs else 0
            return {
                'pairs': pairs, 'risk_matrix': risk_matrix,
                'max_risk': max_risk, 'max_sim': max_sim,
            }

        elif name == 'key_info':
            pairs, risk_matrix = compute_all_pairs(file_data, {
                'text_sim': False, 'key_info': True,
                'file_attr': False, 'image_sim': False
            })
            key_info_matches = build_key_info_matches(pairs)
            return {'key_info_matches': key_info_matches}

        elif name == 'file_attr':
            attr_details = build_attr_details(file_data)
            author_groups = {}
            for ad in attr_details:
                a = ad.get('author', '')
                if a:
                    author_groups[a] = author_groups.get(a, []) + [ad['filename']]
            return {'attr_details': attr_details, 'author_groups': author_groups}

        elif name == 'typo':
            from app.services.typo_detector import detect_typos_batch
            results = detect_typos_batch(file_data)
            return {'results': results}

        elif name == 'relationship':
            from app.services.relationship_extractor import extract_relationships
            report = extract_relationships(file_data)
            return {'report': report}

        elif name == 'quote':
            from app.services.quote_anomaly import compare_bidders_quotes
            result = compare_bidders_quotes(file_data)
            return {'result': result}

    except Exception as e:
        logger.warning(f"Checker {name} failed: {e}")
        return {'error': str(e), 'skipped': True}


def run_analysis(file_data, user_id=None, thread_id=None, tender_text=None,
                 open_info=None, eval_criteria=None):
    """Run all checkers, build a structured report matching the example format.

    Args:
        file_data: list of dicts with {'filename', 'text', 'metadata', 'images'}
        user_id: optional user identifier
        thread_id: optional thread identifier
        tender_text: 招标文件全文（用于 text_sim 模板去除 + 评审标准提取）
        open_info: 结构化开标信息表 dict (from clearance_openinfo.parse_open_info_file)
        eval_criteria: 结构化评审标准 dict (from clearance_openinfo.extract_eval_criteria)

    Returns:
        dict with basic_info, suspected_units, indicators, personnel_summary
    """
    n = len(file_data)
    _ts = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')

    # FIX-015 (C2): if an open-info table provides authoritative bid prices,
    # inject them into file_data BEFORE running checkers so the quote checker
    # uses real totals instead of extract_prices garbage (quote_with_open_price
    # was previously dead code).
    ref_price = None
    if open_info and open_info.get('rows'):
        try:
            from app.services.clearance_openinfo import quote_with_open_price
            file_data, ref_price = quote_with_open_price(file_data, open_info, ref_price)
        except Exception as e:
            logger.warning(f"quote_with_open_price injection failed: {e}")

    # Run all checkers (each in try-except)
    checker_data = {}
    for ind in INDICATOR_DEFS:
        checker_data[ind['checker']] = _run_checker(
            ind['checker'], file_data, user_id, thread_id, tender_text=tender_text
        )

    # OPEN_INFO / TENDER indicators (activate skipped indicators with new inputs)
    open_info_results = {}
    try:
        if open_info or eval_criteria:
            from app.services.clearance_openinfo import compute_open_info_indicators
            open_info_results = compute_open_info_indicators(file_data, open_info, eval_criteria)
    except Exception as e:
        logger.warning(f"OPEN_INFO indicators failed: {e}")

    # Build pairs and risk matrix from text_sim
    text_data = checker_data.get('text_sim', {})
    pairs = text_data.get('pairs', [])
    risk_matrix = text_data.get('risk_matrix', [[0]*n for _ in range(n)])
    max_risk = text_data.get('max_risk', 0)

    # Compute risk scores per file
    file_scores = {i: 0.0 for i in range(n)}
    for p in pairs:
        file_scores[p['i']] += p['risk'] * 0.5
        file_scores[p['j']] += p['risk'] * 0.5

    # Add scores from other checkers
    ki_data = checker_data.get('key_info', {})
    if ki_data.get('key_info_matches'):
        for ki in ki_data['key_info_matches']:
            for i in range(n):
                if file_data[i]['filename'] in (ki['name1'], ki['name2']):
                    file_scores[i] += min(len(ki.get('common_keywords', [])), 10) * 0.5

    attr_data = checker_data.get('file_attr', {})
    author_groups = attr_data.get('author_groups', {})
    for author, files in author_groups.items():
        if len(files) >= 2:
            for i in range(n):
                if file_data[i]['filename'] in files:
                    file_scores[i] += 5.0

    typo_data = checker_data.get('typo', {})
    if typo_data.get('results'):
        for name, report in typo_data['results'].items():
            for i in range(n):
                if file_data[i]['filename'] == name:
                    file_scores[i] += min(report.total_suspects * 0.5, 10)

    rel_data = checker_data.get('relationship', {})
    if rel_data.get('report') and hasattr(rel_data['report'], 'risk_score'):
        file_scores_boost = rel_data['report'].risk_score * 0.1
        for i in range(n):
            file_scores[i] += file_scores_boost

    quote_data = checker_data.get('quote', {})
    if quote_data.get('result'):
        qr = quote_data['result']
        for pb in qr.get('per_bidder', []):
            for i in range(n):
                if file_data[i]['filename'] == pb.get('filename'):
                    file_scores[i] += pb.get('risk_score', 0) * 0.1

    # Boost file scores from OPEN_INFO findings (details reference bidder names)
    for oid, oi in open_info_results.items():
        score = oi.get('score', 0)
        if score <= 0:
            continue
        for d in oi.get('details', []):
            if not isinstance(d, dict):
                continue
            for i in range(n):
                fname = file_data[i]['filename']
                joined = ' '.join(str(v) for v in d.values())
                if fname in joined:
                    file_scores[i] += min(score * 0.1, 5.0)

    # Build suspected units — computed AFTER the indicators loop so the
    # per-file triggered count is a real count of indicators that fired,
    # not a boolean (see FIX-2026-08-28-009).
    suspected_units = []
    for i in range(n):
        score = round(file_scores[i], 1)
        if score > 0:
            suspected_units.append({
                'name': file_data[i]['filename'],
                'indicators_triggered': 0,  # filled after indicators are built
                'score': score,
            })

    # Build indicator results
    indicators = []
    for ind in INDICATOR_DEFS:
        cd = checker_data.get(ind['checker'], {})
        skipped = cd.get('skipped', False)
        error_msg = cd.get('error', '')

        score = 0.0
        result_text = ''
        details = []

        if ind['checker'] == 'text_sim':
            score = max_risk
            high_pairs = [p for p in pairs if p['risk'] > 10]
            bidder_names = set()
            for p in high_pairs:
                bidder_names.add(p['name1'])
                bidder_names.add(p['name2'])
            if max_risk > 10:
                result_text = f"▲ 发现异常：{', '.join(list(bidder_names)[:3])} 等{len(high_pairs)}对组合存在高风险文本相似（最高{max_risk:.1f}），存在围串标嫌疑。"
            else:
                result_text = "√ 未发现显著文本相似异常。"
            # D: 逐对展开到矩阵坐标（所有组合）
            details = [{'pair': f'{p["name1"]} ↔ {p["name2"]}',
                        '矩阵坐标': f'({p.get("i", "-")},{p.get("j", "-")})',
                        'risk': f'{p["risk"]:.1f}', 'sim': f'{p["sim"]:.1f}%'}
                      for p in pairs]
            if skipped:
                result_text = f"○ 跳过（{error_msg}）"

        elif ind['checker'] == 'key_info':
            ki_matches = ki_data.get('key_info_matches', [])
            if ki_matches:
                total = sum(len(ki.get('common_keywords', [])) for ki in ki_matches)
                score = min(total * 0.5, 20)
                result_text = f"▲ 发现{len(ki_matches)}组共享关键信息，共{total}个共同关键词。"
                details = [{'pair': f'{ki["name1"]} ↔ {ki["name2"]}',
                           # FIX-015 (D4): use ki's own coordinates (no more (-,-))
                           '矩阵坐标': f'({ki.get("i", "-")},{ki.get("j", "-")})',
                           'keywords': ', '.join(ki.get('common_keywords', [])[:10])}
                          for ki in ki_matches[:10]]
            else:
                result_text = "√ 未发现重点信息雷同。"
            if skipped:
                result_text = f"○ 跳过（{error_msg}）"

        elif ind['checker'] == 'file_attr':
            group_count = sum(1 for f in author_groups.values() if len(f) >= 2)
            score = group_count * 5.0
            if group_count > 0:
                result_text = f"▲ 发现{group_count}组文件属性雷同。"
                details = [{'author': a, 'files': ', '.join(fs), 'count': len(fs)}
                          for a, fs in author_groups.items() if len(fs) >= 2]
                # D: 补充 pair 级 attr_same
                for p in pairs:
                    if p.get('attr_same'):
                        details.append({'pair': f'{p["name1"]} ↔ {p["name2"]}',
                                        '矩阵坐标': f'({p.get("i", "-")},{p.get("j", "-")})',
                                        'attr_same': '是'})
            else:
                result_text = "√ 未发现文件属性雷同。"
            if skipped:
                result_text = f"○ 跳过（{error_msg}）"

        elif ind['checker'] == 'typo':
            typos = typo_data.get('results', {})
            if typos:
                total = sum(r.total_suspects for r in typos.values())
                critical = sum(r.critical_count for r in typos.values())
                if total > 0:
                    score = min(total * 0.5 + critical * 2, 15)
                    result_text = f"▲ 发现{total}处疑似错别字（其中严重{critical}处）。"
                    details = [{'file': n, 'total': r.total_suspects, 'critical': r.critical_count}
                              for n, r in typos.items() if r.total_suspects > 0][:10]
                else:
                    result_text = "√ 未发现文本质量问题。"
            else:
                result_text = "√ 文本质量检测未发现异常。"
            if skipped:
                result_text = f"○ 跳过（{error_msg}）"

        elif ind['checker'] == 'relationship':
            report = rel_data.get('report')
            # FIX-015 (D3): distinguish the two relationship indicators by what
            # data source they depend on (agent list vs expert list). Without
            # that external data, note the missing source honestly.
            need_external = ind.get('id') in ('bidder_agent_contact', 'expert_bidder_closeness')
            ext_note = ind.get('skip_reason', '')
            if report:
                rs = report.risk_score if hasattr(report, 'risk_score') else 0
                red = len(report.red_flags) if hasattr(report, 'red_flags') else 0
                ents = len(report.entities) if hasattr(report, 'entities') else 0
                rels = len(report.relationships) if hasattr(report, 'relationships') else 0
                score = min(rs, 30)
                if need_external and not ext_note:
                    ext_note = f"（提示：需{ext_note or '外部名单数据'}，当前仅依据投标文件内关系）"
                if red > 0:
                    result_text = f"● 发现{red}个红色预警，{ents}个实体，{rels}个关系，风险评分{rs:.1f}。{ext_note}"
                    details = [{'flag': f} for f in (report.red_flags if hasattr(report, 'red_flags') else [])[:10]]
                else:
                    result_text = f"√ 未发现高危关联，{ents}个实体，{rels}个关系，风险评分{rs:.1f}。{ext_note}"
            else:
                result_text = "√ 未发现关联关系。"
            if skipped:
                result_text = f"○ 跳过（{error_msg}）"

        elif ind['checker'] == 'quote':
            # FIX-015 (C4): high_price_abnormal and quote_proportional_float
            # now judge INDEPENDENTLY (they previously shared the same branch
            # and produced identical output).
            qr = quote_data.get('result', {})
            per_bidder = qr.get('per_bidder', [])
            cross_prog = qr.get('cross_progression', False)
            cross_prog_type = qr.get('cross_progression_type', '')
            is_progression_ind = (ind['id'] == 'quote_proportional_float')

            if not per_bidder:
                result_text = "√ 报价分析未发现异常。"
            elif is_progression_ind:
                # 等比/等差浮动异常：只看报价规律性
                if cross_prog:
                    label = '等差' if cross_prog_type == 'arithmetic' else '等比'
                    score = min(max(qr.get('max_risk_score', 0) + 15, 0), 30)
                    result_text = f"▲ 各投标报价呈{label}规律分布，存在定向陪标嫌疑。最高风险评分{qr.get('max_risk_score', 0):.1f}。"
                    details = [{'bidder': pb.get('filename', ''), 'risk': f'{pb.get("risk_score", 0):.1f}',
                               'cv': f'{pb.get("cv", 0):.4f}',
                               'progression': cross_prog_type or '—'}
                              for pb in per_bidder]
                else:
                    score = 0.0
                    result_text = "√ 各投标报价未见等比/等差规律浮动。"
                    details = []
            else:
                # 高价投标异常：只看报价异常（same_rate/abnormal_drop/clustering/daxie）
                max_qr = qr.get('max_risk_score', 0)
                suspicious = [pb for pb in per_bidder if pb.get('risk_score', 0) > 10]
                score = min(max_qr, 20)
                if suspicious:
                    result_text = f"▲ 发现{len(suspicious)}个投标单位报价疑义。最高风险评分{max_qr:.1f}。"
                    details = [{'bidder': pb.get('filename', ''), 'risk': f'{pb.get("risk_score", 0):.1f}',
                               'cv': f'{pb.get("cv", 0):.4f}',
                               'flags': ', '.join([k for k in ['same_rate_flag', 'abnormal_drop_flag', 'clustering_flag', 'progression_type']
                                                   if pb.get(k)] or ['正常'])}
                              for pb in per_bidder]
                else:
                    score = 0.0
                    result_text = f"√ 报价分析未发现异常。平均CV: {qr.get('avg_cv', 0):.4f}。"
                    details = []
            if skipped:
                result_text = f"○ 跳过（{error_msg}）"

        else:
            # Data-source dependent indicator — skipped placeholder
            skipped = True
            result_text = '○ 需外部数据源（交易平台/评标系统数据）'

        # OPEN_INFO / TENDER indicators: override with real computed results
        oi = open_info_results.get(ind['id'])
        if oi:
            score = oi.get('score', 0)
            result_text = oi.get('result', result_text)
            details = oi.get('details', [])
            skipped = False

        rr = ind.get('rule_ref', {})
        indicators.append({
            'id': ind['id'],
            'name': ind['name'],
            'category': ind['category'],
            'score': round(score, 1),
            'problem': ind['problem'],
            'rule': ind['rule'],
            'rule_ref': rr,
            'rule_text': rule_text(rr.get('law', ''), rr.get('article', '')),
            'result': result_text,
            'details': details,
            'skipped': skipped,
        })

    # Build personnel summary
    personnel = []
    rel_data = checker_data.get('relationship', {})
    report = rel_data.get('report')
    if report and hasattr(report, 'company_personnel_map'):
        cpm = report.company_personnel_map
        for comp in cpm.get('companies', [])[:20]:
            for p in comp.get('personnel', []):
                personnel.append({
                    'company': comp.get('name', ''),
                    'person': p.get('name', ''),
                    'title': p.get('title', ''),
                })

    # 0-100 权重复合指数 (FIX-010)，替代裸加总
    total_score = _weighted_total_score(indicators)

    # Fill per-file indicator trigger counts: an indicator fired for a file if
    # its details reference that file by name, or it's a non-skipped pairwise
    # finding whose pair names include the file.
    for su in suspected_units:
        fname = su['name']
        count = 0
        for ind in indicators:
            if ind.get('skipped'):
                continue
            det = ind.get('details') or []
            referenced = False
            for d in det:
                if isinstance(d, dict):
                    vals = ' '.join(str(v) for v in d.values())
                    if fname in vals:
                        referenced = True
                        break
            if referenced:
                count += 1
        su['indicators_triggered'] = count

    return {
        '_pairs': pairs,
        '_files': [fd.get('filename', '') for fd in file_data],
        'basic_info': {
            'project_name': '用户自定义',
            'bidder_count': n,
            'analysis_date': _ts,
            'total_score': total_score,
            'warning_level': (
                '● 高度预警' if total_score >= 60 else ('◆ 中等预警' if total_score >= 30 else '◇ 正常')
            ),
            # 样本量太小 → 复合指数统计意义有限（P-D 免责声明）
            'sample_note': f'本次共 {n} 家投标人，样本量小，复合指数仅供参考' if n < 5 else '',
        },
        'suspected_units': suspected_units,
        'indicators': indicators,
        'personnel_summary': {
            'total': len(personnel),
            'bidders': sum(1 for p in personnel if '投标' in p.get('title', '')),
            'agents': sum(1 for p in personnel if '代理' in p.get('title', '')),
            'list': personnel[:30],
        },
    }


# ── Word standard font sizes ──
H0 = 42     # 封面标题（初号）
H1 = 16     # 一级标题（三号）
H2 = 15     # 二级/三级标题（小三）
BODY = 12   # 正文/表格内容（小四）
FINE = 10   # 指标详情子段


def _safe(s):
    return str(s) if s else ''


def _add_heading(doc, text, size, bold=True):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    try:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '宋体')
    except Exception:
        pass
    return h


def _add_table(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    return t


def _merge_row(t, row, start, end):
    merged = t.cell(row, start)
    for ci in range(start + 1, end + 1):
        merged = merged.merge(t.cell(row, ci))
    return merged


def _set_cell(cell, text, bold=False, size=BODY):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(_safe(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    try:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '宋体')
    except Exception:
        pass
    return run


def _build_standard_sections(doc, report):
    """Build cover + sections 一~五 (shared by analysis & clearance docx)."""
    info = report['basic_info']

    # ── Cover page ──
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('串通投标线索分析报告')
    run.bold = True
    run.font.size = Pt(H0)

    for _ in range(4):
        doc.add_paragraph()

    doc.add_page_break()

    # ── Section 1: 分析报告综述 ──
    _add_heading(doc, '一、分析报告综述', H1)

    info_table = _add_table(doc, 8, 4)
    bidder_count = info.get('bidder_count', 0)
    total_score = info.get('total_score', 0)
    rows_data = [
        ('标段名称', [info.get('project_name', '用户自定义')] * 3),
        ('标段编号', info.get('bid_number', '—'), '开标时间', info.get('bid_open_time', '—')),
        ('招标单位', info.get('bidder_name', '—'), '招标代理', info.get('agent_name', '—')),
        ('评标办法', info.get('eval_method', '—'), '中标公告发布时间', info.get('award_announce_time', '—')),
        ('中标单位', info.get('winner', '—'), '中标金额', info.get('award_amount', '—')),
        ('冒烟指数', f'{total_score} 分', '预警级别', info.get('warning_level', '')),
        ('地区', info.get('region', '—'), '监督部门', info.get('regulator', '—')),
        ('真实交易平台', [info.get('platform', '—')] * 3),
    ]
    for ri, row in enumerate(rows_data):
        if ri in (0, 7):
            merged = _merge_row(info_table, ri, 1, 3)
            _set_cell(info_table.cell(ri, 0), row[0], bold=True)
            _set_cell(merged, row[1][0])
        else:
            for ci in range(4):
                _set_cell(info_table.cell(ri, ci), row[ci], bold=(ci % 2 == 0))

    doc.add_paragraph()
    # 样本量免责声明（P-D）：投标人 <5 时复合指数统计意义有限
    sample_note = info.get('sample_note', '')
    if sample_note:
        p = doc.add_paragraph()
        rn = p.add_run(sample_note)
        rn.font.size = Pt(FINE)
        rn.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # ── Section 2: 预警单位汇总 ──
    _add_heading(doc, '二、预警单位汇总', H1)
    suspected = report.get('suspected_units', [])
    p = doc.add_paragraph()
    run = p.add_run(f'经过对上述所有指标的深度分析，筛选嫌疑单位{len(suspected)}家(不含招标人与招标代理),详情如下：')
    run.bold = True
    run.font.size = Pt(H2)

    su_table = _add_table(doc, len(suspected) + 1, 2)
    _set_cell(su_table.cell(0, 0), '单位名称', bold=True)
    _set_cell(su_table.cell(0, 1), '涉及指标数量', bold=True)
    for si, su in enumerate(suspected):
        _set_cell(su_table.cell(si + 1, 0), (('★ ' if su['score'] > 10 else '  ') + _safe(su['name'])))
        _set_cell(su_table.cell(si + 1, 1), str(su.get('indicators_triggered', 0)))

    doc.add_paragraph()

    # ── Section 3: 指标详情分析 ──
    _add_heading(doc, '三、指标详情分析', H1)
    p = doc.add_paragraph()
    run = p.add_run(
        '分析指标指的是经过大量串通投标历史案例的调研，分析此类犯罪行为特征、提炼数据规则，形成的'
        '衡量串通投标风险的依据。根据发生串通投标行为的概率为指标进行了归类，分为触发指标、核心指标、'
        '基础指标、扩展指标。'
    )
    run.bold = True
    run.font.size = Pt(H2)

    cat_defs = [
        '触发指标：能够直接作为串通投标线索的指标。',
        '核心指标：能够深层次挖掘围串标主体之间潜在联系的指标。',
        '基础指标：能够从一定程度上反映出串通投标线索的指标。',
        '扩展指标：不影响冒烟指数得分但会引申出其他非主要可疑线索的指标。',
    ]
    for cd in cat_defs:
        p = doc.add_paragraph()
        run = p.add_run(cd)
        run.bold = True
        run.font.size = Pt(BODY)

    indicators = report.get('indicators', [])
    group_meta = [
        ('触发指标', '3.1、'),
        ('核心指标', '3.2、'),
        ('基础指标', '3.3、'),
        ('扩展指标', '3.4、'),
    ]
    for gi, (cat, prefix) in enumerate(group_meta):
        group = [ind for ind in indicators if ind['category'] == cat]
        if not group:
            continue
        _add_heading(doc, f'{prefix}{cat}', H2)
        for si, ind in enumerate(group):
            sub_no = f'3.{gi + 1}.{si + 1}'
            _add_heading(doc, f'{sub_no}、{ind["name"]}', H2)

            t = _add_table(doc, 6, 4)
            _set_cell(t.cell(0, 0), '指标名称', bold=True)
            _merge_row(t, 0, 1, 3)
            _set_cell(t.cell(0, 1), ind['name'])

            _set_cell(t.cell(1, 0), '指标类别', bold=True)
            _set_cell(t.cell(1, 1), ind['category'])
            _set_cell(t.cell(1, 2), '指标得分', bold=True)
            _set_cell(t.cell(1, 3), f"{ind.get('score', 0)} 分")

            _set_cell(t.cell(2, 0), '针对问题', bold=True)
            _merge_row(t, 2, 1, 3)
            _set_cell(t.cell(2, 1), ind.get('problem', ''))

            _set_cell(t.cell(3, 0), '分析规则', bold=True)
            _merge_row(t, 3, 1, 3)
            rule_text_val = ind.get('rule', '')
            rule_ref = ind.get('rule_ref', {})
            law_text = ind.get('rule_text', '')
            if law_text:
                rule_text_val += f"\n规则依据：{rule_ref.get('label', '')} {law_text}"
            _set_cell(t.cell(3, 1), rule_text_val)

            _set_cell(t.cell(4, 0), '分析结果', bold=True)
            _merge_row(t, 4, 1, 3)
            result_run = _set_cell(t.cell(4, 1), ind.get('result', ''))
            result_text = _safe(ind.get('result', ''))
            if '▲' in result_text or '●' in result_text:
                result_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
            elif '√' in result_text:
                result_run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
            elif '○' in result_text:
                result_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

            _set_cell(t.cell(5, 0), '指标详情', bold=True)
            details = ind.get('details', [])
            if details and not ind.get('skipped', False):
                _merge_row(t, 5, 1, 3)
                detail_cell = t.cell(5, 1)
                detail_cell.text = ''
                dp = detail_cell.paragraphs[0]
                if isinstance(details[0], dict):
                    for d in details[:10]:
                        rp = dp
                        for k, v in list(d.items())[:4]:
                            rp = detail_cell.add_paragraph()
                            rn = rp.add_run(f'{k}: {_safe(v)}')
                            rn.font.size = Pt(FINE)
                        if len(details) > 10:
                            rp = detail_cell.add_paragraph()
                            rn = rp.add_run(f'…（共{len(details)}项）')
                            rn.font.size = Pt(FINE)
                            break
                else:
                    rn = dp.add_run(_safe(details[0]))
                    rn.font.size = Pt(FINE)
            else:
                _merge_row(t, 5, 1, 3)

    # ── Section 4: 关系人员汇总 ──
    doc.add_page_break()
    _add_heading(doc, '四、关系人员汇总', H1, bold=True)
    personnel = report.get('personnel_summary', {})
    personnel_list = personnel.get('list', [])
    p = doc.add_paragraph()
    run = p.add_run(
        f'以冒烟嫌疑单位为分析对象，筛选单位项目所属关系人员共计{len(personnel_list)}人，'
        f'其中所属投标人{personnel.get("bidders", 0)}人，所属招标人0人，所属招标代理0人。详情如下：'
    )
    run.bold = True
    run.font.size = Pt(H2)

    pt = _add_table(doc, len(personnel_list) + 1, 4)
    # FIX-015 (D4): drop 身份证号/联系电话 empty shell columns (never populated)
    for j, hdr in enumerate(['单位名称', '姓名', '职务', '人员类型']):
        _set_cell(pt.cell(0, j), hdr, bold=True)
    for pi, pe in enumerate(personnel_list):
        _set_cell(pt.cell(pi + 1, 0), _safe(pe.get('company', '')))
        _set_cell(pt.cell(pi + 1, 1), _safe(pe.get('person', '')))
        _set_cell(pt.cell(pi + 1, 2), _safe(pe.get('title', '')))
        _set_cell(pt.cell(pi + 1, 3), '投标人')

    doc.add_paragraph()

    # ── Section 5: 开标信息表 ──
    _add_heading(doc, '五、开标信息表', H1, bold=True)
    files = report.get('_files', [])
    open_info = report.get('open_info') or {}
    open_rows = open_info.get('rows', []) if isinstance(open_info, dict) else []
    # 用开标表行数，若缺失则退化为投标文件数
    row_count = max(len(open_rows), len(files))
    bid_table = _add_table(doc, row_count + 1, 15)
    headers = ['序号', '开标时间', '投标单位', '联系人', '联系电话', '文件下载IP', '标书上传时间',
               '标书上传IP', '解密状态', '解密IP', '文件码', '加密锁', '报价方式', '投标报价', '备注']
    for j, hdr in enumerate(headers):
        _set_cell(bid_table.cell(0, j), hdr, bold=True, size=8)

    def _fmt_price(v):
        if v is None or v == '':
            return ''
        try:
            num = float(str(v).replace(',', ''))
            return f'{num / 10000:.2f}万元' if abs(num) >= 10000 else f'{num:.2f}元'
        except (ValueError, TypeError):
            return _safe(str(v))

    for fi in range(row_count):
        oi = open_rows[fi] if fi < len(open_rows) else {}
        fname = str(oi.get('bidder', '') or (files[fi] if fi < len(files) else ''))
        cells = {
            0: str(fi + 1),
            1: _safe(oi.get('open_time', '')),
            2: _safe(fname),
            3: _safe(oi.get('contact', '')),
            4: _safe(oi.get('phone', '')),
            12: _safe(oi.get('price_mode', '')),
            13: _fmt_price(oi.get('bid_price')),
            14: _safe(oi.get('remark', '')),
        }
        for j in range(15):
            _set_cell(bid_table.cell(fi + 1, j), cells.get(j, ''), size=8)
    if not open_rows and files:
        note = doc.add_paragraph()
        rn = note.add_run('（开标时间/联系人/报价等数据需导入开标信息表；IP/文件码/加密锁为交易平台字段，当前系统不采集）')
        rn.font.size = Pt(FINE)
        rn.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.add_paragraph()


def _add_date_line(doc, report):
    """Append the closing date line (ends the document)."""
    info = report.get('basic_info', {})
    p = doc.add_paragraph()
    run = p.add_run(_safe(info.get('analysis_date', '')))
    run.font.size = Pt(BODY)


def _new_docx():
    from docx import Document as DocxDocument
    doc = DocxDocument()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(BODY)
    # For east Asian fonts
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return doc


def build_analysis_docx(report, info_overrides=None):
    """Generate .docx following the bid-rigging clue analysis format."""
    if info_overrides:
        report = dict(report)
        report['basic_info'] = {**report.get('basic_info', {}), **info_overrides}
    doc = _new_docx()
    _build_standard_sections(doc, report)
    _add_date_line(doc, report)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── 清标 docx：标准五章 + 六横向对比 / 七合规 / 八 AI 评审 ──

def build_clearance_docx(report, info_overrides=None):
    """Generate the unified 清标 .docx.
    report: dict from clearance_engine.run_clearance()
    """
    if info_overrides:
        report = dict(report)
        report['basic_info'] = {**report.get('basic_info', {}), **info_overrides}
    doc = _new_docx()
    _build_standard_sections(doc, report)

    # ── Section 6: 横向对比分析 ──
    cross = report.get('cross_comparison')
    if cross:
        pairs = cross.get('pairs', [])
        risk_matrix = cross.get('risk_matrix', [])
        text_matrix = cross.get('text_matrix') or risk_matrix
        key_matrix = cross.get('key_matrix') or risk_matrix
        attr_matrix = cross.get('attr_matrix') or risk_matrix
        key_info = cross.get('key_info_matches', [])
        attr = cross.get('attr_details', [])
        gangs = cross.get('gangs', [])
        files = report.get('_files', [])

        def _write_matrix(title, m, fmt, extra=''):
            _add_heading(doc, title, H2)
            n = len(files)
            t = _add_table(doc, n + 1, n + 1)
            _set_cell(t.cell(0, 0), '投标单位', bold=True)
            for j, fname in enumerate(files):
                _set_cell(t.cell(0, j + 1), _safe(truncate_filename(fname, 8)), bold=True)
            for i in range(n):
                _set_cell(t.cell(i + 1, 0), _safe(truncate_filename(files[i], 8)), bold=True)
                for j in range(n):
                    val = '--' if i == j else (
                        fmt(m[i][j]) if i < len(m) and j < len(m[i]) else '--')
                    _set_cell(t.cell(i + 1, j + 1), val)
            doc.add_paragraph()

        _add_heading(doc, '六、横向对比分析', H1, bold=True)

        # 6.1 综合风险矩阵
        _write_matrix('6.1、综合风险矩阵', risk_matrix, lambda v: f"{v:.1f}")

        # E1: 多维矩阵（文本/关键信息/属性）
        _write_matrix('6.2、文本相似度矩阵（%）', text_matrix, lambda v: f"{v:.1f}")
        _write_matrix('6.3、关键信息重合矩阵（%）', key_matrix, lambda v: f"{v:.1f}")
        _write_matrix('6.4、文件属性雷同矩阵（1/0）', attr_matrix, lambda v: f"{int(round(v))}")

        # 6.5 高风险对详情
        high_pairs = [p for p in pairs if p.get('risk', 0) > 5]
        if high_pairs:
            _add_heading(doc, '6.5、高风险组合详情', H2)
            pt = _add_table(doc, len(high_pairs) + 1, 7)
            for j, hdr in enumerate(['投标单位1', '投标单位2', '风险度', '文本相似度', '关键信息重合', '文件属性雷同', '矩阵坐标']):
                _set_cell(pt.cell(0, j), hdr, bold=True, size=9)
            for ri, p in enumerate(high_pairs):
                _set_cell(pt.cell(ri + 1, 0), _safe(p.get('name1', '')), size=9)
                _set_cell(pt.cell(ri + 1, 1), _safe(p.get('name2', '')), size=9)
                _set_cell(pt.cell(ri + 1, 2), f"{p.get('risk', 0):.1f}", size=9)
                _set_cell(pt.cell(ri + 1, 3), f"{p.get('sim', 0):.1f}%", size=9)
                _set_cell(pt.cell(ri + 1, 4), f"{p.get('key_sim', 0):.1f}%", size=9)
                _set_cell(pt.cell(ri + 1, 5), '是' if p.get('attr_same') else '否', size=9)
                _set_cell(pt.cell(ri + 1, 6), f"({p.get('i', '-')},{p.get('j', '-')})", size=9)
            doc.add_paragraph()

        # 6.6 全部组合明细（C：不设阈值）
        if pairs:
            _add_heading(doc, '6.6、全部组合明细', H2)
            at = _add_table(doc, len(pairs) + 1, 7)
            for j, hdr in enumerate(['投标单位1', '投标单位2', '风险度', '文本相似度', '关键信息重合', '文件属性雷同', '矩阵坐标']):
                _set_cell(at.cell(0, j), hdr, bold=True, size=9)
            has_mismatch = any(p.get('component_mismatch') for p in pairs)
            for ri, p in enumerate(pairs):
                _set_cell(at.cell(ri + 1, 0), _safe(p.get('name1', '')), size=9)
                _set_cell(at.cell(ri + 1, 1), _safe(p.get('name2', '')), size=9)
                _set_cell(at.cell(ri + 1, 2), f"{p.get('risk', 0):.1f}", size=9)
                sim_txt = f"{p.get('sim', 0):.1f}%"
                if p.get('component_mismatch'):
                    sim_txt = '异组件(不计)'
                _set_cell(at.cell(ri + 1, 3), sim_txt, size=9)
                _set_cell(at.cell(ri + 1, 4), f"{p.get('key_sim', 0):.1f}%", size=9)
                _set_cell(at.cell(ri + 1, 5), '是' if p.get('attr_same') else '否', size=9)
                _set_cell(at.cell(ri + 1, 6), f"({p.get('i', '-')},{p.get('j', '-')})", size=9)
            if has_mismatch:
                note = doc.add_paragraph()
                rn = note.add_run('注：【异组件】比对（如价格标↔技术标）的结构性条款重叠不作串标依据，文本相似度已不计入风险。')
                rn.font.size = Pt(FINE)
                rn.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
            doc.add_paragraph()

        # 6c. 关键信息匹配表
        if key_info:
            _add_heading(doc, '6.7、重点信息雷同', H2)
            kt = _add_table(doc, len(key_info) + 1, 3)
            for j, hdr in enumerate(['投标单位1', '投标单位2', '共同关键词']):
                _set_cell(kt.cell(0, j), hdr, bold=True, size=8)
            for ri, ki in enumerate(key_info):
                _set_cell(kt.cell(ri + 1, 0), _safe(ki.get('name1', '')), size=8)
                _set_cell(kt.cell(ri + 1, 1), _safe(ki.get('name2', '')), size=8)
                _set_cell(kt.cell(ri + 1, 2), ', '.join(ki.get('common_keywords', [])[:10]), size=8)
            doc.add_paragraph()

        # E2: 疑似围标集团
        if gangs:
            _add_heading(doc, '6.8、疑似围标集团', H2)
            gt = _add_table(doc, len(gangs) + 1, 5)
            for j, hdr in enumerate(['集团', '成员单位', '成员数', '组内最高风险', '组内平均风险']):
                _set_cell(gt.cell(0, j), hdr, bold=True, size=9)
            for ri, g in enumerate(gangs):
                _set_cell(gt.cell(ri + 1, 0), f"集团{ri + 1}", size=9)
                _set_cell(gt.cell(ri + 1, 1), ' / '.join(_safe(f) for f in g.get('files', [])), size=9)
                _set_cell(gt.cell(ri + 1, 2), str(len(g.get('members', []))), size=9)
                _set_cell(gt.cell(ri + 1, 3), f"{g.get('max_risk', 0):.1f}", size=9)
                _set_cell(gt.cell(ri + 1, 4), f"{g.get('avg_risk', 0):.1f}", size=9)
            doc.add_paragraph()

    # ── Section 7: 合规审查结果 ──
    comp = report.get('compliance')
    if comp and not comp.get('skipped'):
        _add_heading(doc, '七、合规审查结果', H1, bold=True)
        p = doc.add_paragraph()
        run = p.add_run(f'基于招标文件《{_safe(comp.get("tender_name", ""))}》提取的 {len(comp.get("rules", []))} 条规则，对各单位投标文件逐一审查。')
        run.bold = True
        run.font.size = Pt(BODY)

        for pf in comp.get('per_file', []):
            _add_heading(doc, f'7.{comp.get("per_file", []).index(pf) + 1}、{_safe(pf.get("filename", ""))}', H2)
            s = pf.get('summary', {})
            st = _add_table(doc, 2, 4)
            for j, hdr in enumerate(['通过', '警告', '违规', '严重']):
                _set_cell(st.cell(0, j), hdr, bold=True, size=9)
            for j, key in enumerate(['pass', 'warning', 'violation', 'critical']):
                _set_cell(st.cell(1, j), str(s.get(key, 0)), size=9)
            doc.add_paragraph()

            results = pf.get('results', [])
            if results:
                rule_map = {r.get('rule_id'): r.get('description', '') for r in comp.get('rules', [])}
                rt = _add_table(doc, len(results) + 1, 5)
                for j, hdr in enumerate(['规则编号', '规则描述', '审查结论', '证据', '建议']):
                    _set_cell(rt.cell(0, j), hdr, bold=True, size=8)
                for ri, res in enumerate(results):
                    _set_cell(rt.cell(ri + 1, 0), _safe(res.get('rule_id', '')), size=8)
                    _set_cell(rt.cell(ri + 1, 1), _safe(rule_map.get(res.get('rule_id', ''), ''))[:80], size=8)
                    _set_cell(rt.cell(ri + 1, 2), _safe(res.get('verdict', '')), size=8)
                    _set_cell(rt.cell(ri + 1, 3), _safe(res.get('evidence', ''))[:80], size=8)
                    _set_cell(rt.cell(ri + 1, 4), _safe(res.get('suggestion', ''))[:80], size=8)
                doc.add_paragraph()

    # ── Section 8: AI 综合评审 ──
    ai = report.get('ai_review')
    if ai and ai.get('per_file'):
        _add_heading(doc, '八、AI 综合评审', H1, bold=True)
        for pi, pf in enumerate(ai.get('per_file', [])):
            review = pf.get('review', {})
            _add_heading(doc, f'8.{pi + 1}、{_safe(pf.get("filename", ""))}', H2)

            scores = review.get('scores') or {}
            if scores:
                st = _add_table(doc, 2, len(scores) + 2)
                _set_cell(st.cell(0, 0), '综合评分', bold=True, size=9)
                _set_cell(st.cell(1, 0), _safe(review.get('overall', '')), size=9)
                for j, (k, v) in enumerate(scores.items(), start=1):
                    _set_cell(st.cell(0, j), _safe(k), bold=True, size=9)
                    _set_cell(st.cell(1, j), _safe(v), size=9)
                _set_cell(st.cell(0, len(scores) + 1), '结论', bold=True, size=9)
                _set_cell(st.cell(1, len(scores) + 1), _safe(review.get('verdict', '')), size=9)
                doc.add_paragraph()

            issues = review.get('issues') or []
            if issues:
                it = _add_table(doc, len(issues) + 1, 4)
                for j, hdr in enumerate(['维度', '严重度', '问题', '建议']):
                    _set_cell(it.cell(0, j), hdr, bold=True, size=8)
                for ri, iss in enumerate(issues):
                    _set_cell(it.cell(ri + 1, 0), _safe(iss.get('axis', '')), size=8)
                    _set_cell(it.cell(ri + 1, 1), _safe(iss.get('severity', '')), size=8)
                    _set_cell(it.cell(ri + 1, 2), _safe(iss.get('finding', ''))[:100], size=8)
                    _set_cell(it.cell(ri + 1, 3), _safe(iss.get('suggestion', ''))[:100], size=8)
                doc.add_paragraph()

            if review.get('summary'):
                p = doc.add_paragraph()
                run = p.add_run(f"AI 综合意见：{_safe(review.get('summary', ''))}")
                run.font.size = Pt(BODY)

    # ── Section 9: 图片随机抽检说明 ──
    image_sampling = report.get('image_sampling') or []
    if image_sampling:
        _add_heading(doc, '九、图片随机抽检说明', H1, bold=True)
        p = doc.add_paragraph()
        run = p.add_run('对投标文件中的图片采用随机抽检方式识别，每文件随机抽取最多 20 张由视觉模型描述，'
                        '并记录所在位置及前后文以便复核。')
        run.bold = True
        run.font.size = Pt(BODY)

        for entry in image_sampling:
            fname = _safe(entry.get('filename', ''))
            samples = entry.get('samples', [])
            _add_heading(doc, f'9.{image_sampling.index(entry) + 1}、{fname}', H2)
            if not samples:
                p = doc.add_paragraph()
                run = p.add_run('该文件未抽检到可识别图片，或视觉模型已熔断跳过。')
                run.font.size = Pt(BODY)
                continue
            st = _add_table(doc, len(samples) + 1, 5)
            for j, hdr in enumerate(['抽检编号', '所在位置', '前文(10字)', '后文(10字)', '识别摘要']):
                _set_cell(st.cell(0, j), hdr, bold=True)
            for ri, s in enumerate(samples):
                _set_cell(st.cell(ri + 1, 0), str(s.get('seq', ri + 1)))
                _set_cell(st.cell(ri + 1, 1), _safe(s.get('chapter', ''))[:40])
                _set_cell(st.cell(ri + 1, 2), _safe(s.get('prev', '')))
                _set_cell(st.cell(ri + 1, 3), _safe(s.get('next', '')))
                _set_cell(st.cell(ri + 1, 4), _safe(s.get('desc', ''))[:50])
            doc.add_paragraph()

    # ── Section 10: 全量审计补充检查 ──
    audit_supplement = report.get('audit_supplement')
    if audit_supplement and audit_supplement.get('per_file'):
        _add_heading(doc, '十、全量审计补充检查', H1, bold=True)
        p = doc.add_paragraph()
        run = p.add_run('基于全量审计函数（风格分析、规则提取）对各投标文件进行补充检查，'
                        '结果纳入清标综合评估。时间线合规因上传场景无项目时间线数据而占位跳过。')
        run.bold = True
        run.font.size = Pt(BODY)

        for entry in audit_supplement['per_file']:
            fname = _safe(entry.get('filename', ''))
            style = entry.get('style') or {}
            rules = entry.get('rules') or {}
            timeline = entry.get('timeline') or {}

            _add_heading(doc, fname or '—', H2)

            # 风格分析表
            st = _add_table(doc, 2, 3)
            _set_cell(st.cell(0, 0), '风格分析', bold=True)
            _set_cell(st.cell(0, 1), '评分', bold=True)
            _set_cell(st.cell(0, 2), '等级', bold=True)
            _set_cell(st.cell(1, 0), '正式性')
            _set_cell(st.cell(1, 1), f"{style.get('score', '—')}")
            _set_cell(st.cell(1, 2), _safe(style.get('findings', {}).get('formality_label', '—')))
            doc.add_paragraph()

            # 规则提取表
            rt = _add_table(doc, 2, 3)
            _set_cell(rt.cell(0, 0), '自规则提取', bold=True)
            _set_cell(rt.cell(0, 1), '评分', bold=True)
            _set_cell(rt.cell(0, 2), '提取条数', bold=True)
            _set_cell(rt.cell(1, 0), '规则数量')
            _set_cell(rt.cell(1, 1), f"{rules.get('score', '—')}")
            _set_cell(rt.cell(1, 2), f"{rules.get('count', 0)}")
            doc.add_paragraph()

            # 时间线合规（占位）
            p = doc.add_paragraph()
            run = p.add_run('时间线合规：')
            run.bold = True
            run.font.size = Pt(BODY)
            note = _safe(timeline.get('note', ''))
            run2 = p.add_run(note or '无项目时间线数据，跳过')
            run2.font.size = Pt(BODY)
            doc.add_paragraph()

    _add_date_line(doc, report)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── LibreOffice PDF conversion ──

def convert_docx_to_pdf(docx_bytes, task_id=''):
    """Convert .docx bytes to .pdf bytes via headless LibreOffice.

    Uses LIBREOFFICE_BIN env (default: 'soffice'). Runs in a temp dir so
    --convert-to keeps a clean output filename. Returns None if the
    soffice binary is unavailable (caller falls back to DOCX-only ZIP).

    Args:
        docx_bytes: bytes of the .docx document
        task_id: optional id used in the temp file name (avoid collisions)

    Returns:
        bytes of the converted .pdf, or None if LibreOffice unavailable
    """
    import os as _os, shutil, subprocess, tempfile, time as _time

    soffice = _os.environ.get('LIBREOFFICE_BIN', 'soffice')
    if not shutil.which(soffice):
        return None

    with tempfile.TemporaryDirectory(prefix='docx2pdf_') as tmp:
        src_name = f"report_{task_id or int(_time.time())}.docx"
        src = _os.path.join(tmp, src_name)
        with open(src, 'wb') as f:
            f.write(docx_bytes)
        result = subprocess.run(
            [soffice, '--headless', '--convert-to', 'pdf', '--outdir', tmp, src],
            capture_output=True, timeout=120,
        )
        if result.returncode != 0:
            logger.warning('soffice convert failed: %s', result.stderr.decode('utf-8', 'replace')[:500])
            return None
        pdf_path = _os.path.join(tmp, src_name.replace('.docx', '.pdf'))
        if not _os.path.exists(pdf_path):
            return None
        with open(pdf_path, 'rb') as f:
            return f.read()


# ── Celery async task ──

@celery_app.task(bind=True, name='document_analysis_task', max_retries=1,
                 soft_time_limit=2400, time_limit=2700)
def run_analysis_async(self, file_data, file_specs, user_id, thread_id, task_id, project_id=None):
    """Celery task: runs full analysis + generates DOCX + stores ZIP in DB.

    file_data:  pre-extracted dicts (legacy small uploads)
    file_specs: [{'abs_path','filename'}] — extracted here in the worker,
                page-by-page (large-file friendly).
    """
    from app.services.task_bus import TaskBus
    from app.services.file_processing import extract_text_from_path, extract_metadata_from_path
    import os as _os, zipfile, hashlib, json

    # Extraction helpers touch flask.session (analyze_images pref) which
    # needs a *request* context — test_request_context provides one for the
    # whole task.
    from celery_app import init_flask_context
    _flask_app = init_flask_context()
    _req_ctx = _flask_app.test_request_context()
    _req_ctx.push()

    bus = TaskBus(task_id, 'doc_analysis', '投标文档深度分析')
    bus.start()

    try:
        # Step 1: Run analysis (25%)
        bus.progress(5, '正在提取文件内容...')

        # Worker-side extraction for pre-uploaded files (bounded memory)
        all_file_data = list(file_data or [])
        for i, spec in enumerate(file_specs or []):
            fname = spec.get('filename', '')
            bus.progress(5 + int(15 * i / max(1, len(file_specs))),
                         f'正在提取文件内容 ({i + 1}/{len(file_specs)}): {fname}...')
            text, _ = extract_text_from_path(spec['abs_path'], fname)
            if not text or text.startswith("["):
                logger.warning(f"Skipping unreadable file: {fname}")
                continue
            meta = extract_metadata_from_path(spec['abs_path'], fname)
            all_file_data.append({
                'filename': fname,
                'text': text,
                'metadata': meta or {},
                'images': [],
            })

        if not all_file_data:
            bus.fail('无可提取文本的文件，任务终止')
            return

        n = len(all_file_data)

        bus.progress(20, f'开始分析 {n} 个文件...')

        # Step 2: Run all checkers
        bus.progress(25, '正在进行文本相似度分析...')
        report = run_analysis(all_file_data, user_id, thread_id)

        bus.progress(40, '正在进行重点信息雷同分析...')
        bus.progress(55, '正在进行文件属性对比...')
        bus.progress(65, '正在进行文本质量检测...')
        bus.progress(75, '正在进行关联关系分析...')
        bus.progress(85, '正在进行报价异常检测...')

        # Step 3: Generate DOCX + PDF (90%)
        bus.progress(90, '正在生成分析报告...')

        # Initialize Flask context for DB access
        from celery_app import init_flask_context
        init_flask_context()

        from app.config import DATA_DIR, to_rel_path
        docx_bytes = build_analysis_docx(report)

        # Step 4: Convert to PDF via LibreOffice (headless, Docker). Falls back to DOCX-only.
        pdf_bytes = None
        try:
            pdf_bytes = convert_docx_to_pdf(docx_bytes, task_id)
        except Exception as _e:
            logger.warning('PDF conversion failed, falling back to DOCX only: %s', _e)
            bus.progress(95, 'PDF 转换不可用，仅输出 DOCX')
            pdf_bytes = None

        # Store ZIP in DB
        batch_dir = _os.path.join(DATA_DIR, 'batch_results')
        _os.makedirs(batch_dir, exist_ok=True)
        zip_name = f"doc_analysis_{task_id}.zip"
        zip_path = _os.path.join(batch_dir, zip_name)
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f"串通投标线索分析报告_{task_id[:8]}.docx", docx_bytes)
            if pdf_bytes:
                zf.writestr(f"串通投标线索分析报告_{task_id[:8]}.pdf", pdf_bytes)

        file_names = [fd['filename'] for fd in all_file_data]
        from app.database import get_db_connection
        with get_db_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    INSERT INTO batch_comparison_results (user_id, task_id, project_id, file_count, pair_count, max_risk, file_names, zip_path)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    user_id, task_id, project_id, len(all_file_data), 0,
                    report['basic_info'].get('total_score', 0),
                    json.dumps(file_names, ensure_ascii=False),
                    to_rel_path(zip_path),
                ))

                pairs = report.get('_pairs', [])
                if pairs:
                    ranked = sorted(pairs, key=lambda p: p.get('risk', 0), reverse=True)
                    for rank, p in enumerate(ranked):
                        risk_scores = {
                            'text_sim': round(p.get('sim', 0), 2),
                            'attr_same': p.get('attr_same', 0),
                            'risk': round(p.get('risk', 0), 2),
                        }
                        cur.execute("""
                            INSERT INTO batch_pair_results (task_id, file_a, file_b, similarity, max_risk, risk_scores, pair_rank)
                            VALUES (%s, %s, %s, %s, %s, %s, %s)
                            ON CONFLICT (task_id, file_a, file_b) DO UPDATE SET
                                similarity = EXCLUDED.similarity,
                                max_risk = EXCLUDED.max_risk,
                                risk_scores = EXCLUDED.risk_scores,
                                pair_rank = EXCLUDED.pair_rank
                        """, (
                            task_id,
                            p.get('name1', ''),
                            p.get('name2', ''),
                            round(p.get('sim', 0), 2),
                            round(p.get('risk', 0), 2),
                            json.dumps(risk_scores, ensure_ascii=False),
                            rank + 1,
                        ))

                conn.commit()

        from flask import url_for
        try:
            download_url = url_for('batch.download_batch_result', task_id=task_id)
        except Exception:
            from app.config import BASE_DIR
            download_url = f'/batch_result/{task_id}'

        bus.progress(100, '分析完成')
        bus.complete({
            'report': report,
            'download_url': download_url,
            'file_count': len(all_file_data),
        })

    except Exception as e:
        logger.error(f"Analysis task failed: {e}", exc_info=True)
        import traceback as _tb
        _frames = _tb.format_exc().strip().split('\n')
        bus.fail(f"{e} | at: {_frames[-1].strip()}"[:500])
