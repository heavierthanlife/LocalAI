"""Template Renderer — generates .docx documents from 模板 (template) skills.
    Parses structured markdown from a template skill and produces
    a formatted Word document using the existing file_generator infrastructure."""

import logging
from io import BytesIO

logger = logging.getLogger(__name__)


def _parse_sections(skill_md: str) -> dict:
    """Parse a skill markdown into sections: {header: [items]}.
    
    Items are bullet points (- ...) or numbered steps (1. ...) under each ## section.
    """
    sections = {}
    current_section = None
    for line in skill_md.split('\n'):
        if line.startswith('## '):
            current_section = line[3:].strip()
            sections[current_section] = []
        elif current_section is not None:
            stripped = line.strip()
            if stripped.startswith('- '):
                sections[current_section].append(stripped[2:])
            elif stripped and stripped[0].isdigit() and '. ' in stripped[:4]:
                sections[current_section].append(stripped.split('. ', 1)[1] if '. ' in stripped else stripped)
    return sections


def render_template_to_markdown(skill_md: str, title: str = None) -> str:
    """Convert a template skill markdown into a clean output markdown suitable for docx conversion.
    
    Strips the skill header/metadata, keeps only the structured content.
    """
    sections = _parse_sections(skill_md)
    if not sections:
        return skill_md

    lines = []
    if title:
        lines.append(f"# {title}")
        lines.append("")

    for header, items in sections.items():
        if not items:
            continue
        lines.append(f"## {header}")
        for item in items:
            lines.append(f"- {item}")
        lines.append("")

    return '\n'.join(lines)


def render_template_to_docx(skill_md: str, title: str = None) -> bytes:
    """Generate a .docx from a template skill's markdown structure.
    
    Returns bytes of the .docx file.
    Uses file_generator.markdown_to_docx for the heavy lifting.
    """
    clean_md = render_template_to_markdown(skill_md, title=title)
    try:
        from app.services.file_generator import markdown_to_docx
        return markdown_to_docx(clean_md, title or '文档')
    except Exception as e:
        logger.warning(f"file_generator.markdown_to_docx failed: {e}, trying fallback")
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            doc = Document()
            section = doc.sections[0]
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

            if title:
                p = doc.add_heading(title, level=0)
                for run in p.runs:
                    run.font.size = Pt(16)

            sections = _parse_sections(clean_md)
            for header, items in sections.items():
                if not items:
                    continue
                doc.add_heading(header, level=1)
                for item in items:
                    doc.add_paragraph(item, style='List Bullet')

            buf = BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception as e2:
            logger.error(f"Fallback docx generation also failed: {e2}")
            raise
