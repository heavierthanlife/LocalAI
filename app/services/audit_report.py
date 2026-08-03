"""Unified Bid Audit Report Generator — DOCX + XLSX with TOC and scoring."""
import json
import logging
import os
from datetime import datetime, timezone

from app.config import DATA_DIR
from app.database import get_db_connection

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

logger = logging.getLogger(__name__)

AUDIT_REPORTS_DIR = DATA_DIR / 'audit_reports'

FUNCTION_LABELS = {
    'rule_extraction': 'Rule Extraction',
    'compliance_check': 'Compliance Check',
    'typo_detection': 'Typo Detection',
    'quote_anomaly': 'Quote Anomaly',
    'relationship_extraction': 'Relationship Analysis',
    'ai_doc_review': 'AI Document Review',
    'style_analysis': 'Style Analysis',
}

FUNCTION_LABELS_ZH = {
    'rule_extraction': '规则提取',
    'compliance_check': '合规审查',
    'typo_detection': '错别字检测',
    'quote_anomaly': '报价异常',
    'relationship_extraction': '关系分析',
    'ai_doc_review': 'AI文档审查',
    'style_analysis': '文风分析',
}


def _load_run_data(run_id: int) -> dict | None:
    from app.services.audit_engine import get_run_results
    return get_run_results(run_id)


def _get_scoring_rationale(func_name: str, findings: dict,
                            text_length: int | None = None) -> dict:
    """Return scoring rationale with base value, deductions, and formula.

    Returns:
        {'base_value': float, 'actual_score': float, 'deductions': list[dict],
         'formula_zh': str, 'detail_zh': str}
    """
    base = 100.0
    deductions = []
    formula = ''
    detail = ''

    if func_name == 'rule_extraction':
        rules = findings.get('rules', [])
        found = len(rules) if isinstance(rules, list) else 0
        from app.services.audit_engine import SCORING_FUNCTIONS
        actual = SCORING_FUNCTIONS[func_name](findings)
        formula = f'基准分{base}分 × (提取规则数{found} / 预期最少规则数)'
        if found == 0:
            deductions.append({'reason': '未提取到任何规则', 'amount': 100})
        elif actual < base:
            short = base - actual
            deductions.append({'reason': f'提取规则数({found}条)未达预期', 'amount': round(short, 1)})

    elif func_name == 'compliance_check':
        results = findings.get('results', [])
        s = findings.get('summary', {})
        total = len(results) if results else 0
        passed = s.get('pass', 0)
        warnings = s.get('warning', 0)
        violations = s.get('violation', 0)
        critical = s.get('critical', 0)
        from app.services.audit_engine import SCORING_FUNCTIONS
        actual = SCORING_FUNCTIONS[func_name](findings)
        formula = f'基准分{base}分 × (通过项{passed} / 总检查项{total})'
        fail_count = warnings + violations + critical
        if fail_count > 0:
            deductions.append({'reason': f'警告{warnings}项、违规{violations}项、严重{critical}项', 'amount': round(base - actual, 1)})

    elif func_name == 'typo_detection':
        flist = findings.get('findings', [])
        count = len(flist) if isinstance(flist, list) else findings.get('total_findings', 0)
        from app.services.audit_engine import SCORING_FUNCTIONS
        actual = SCORING_FUNCTIONS[func_name](findings, text_length or 500)
        formula = f'基准分{base}分 - (每万字错别字数 × 扣分权重)'
        if count > 0:
            deductions.append({'reason': f'发现{count}个错别字', 'amount': round(base - actual, 1)})

    elif func_name == 'quote_anomaly':
        sev = findings.get('severity_index', 0)
        from app.services.audit_engine import SCORING_FUNCTIONS
        actual = SCORING_FUNCTIONS[func_name](findings)
        formula = f'基准分{base}分 - 异常严重度指数({sev})'
        if sev > 0:
            deductions.append({'reason': f'异常严重度指数 {sev}', 'amount': round(min(base, sev), 1)})

    elif func_name == 'relationship_extraction':
        signals = findings.get('collusion_signals', [])
        red_flags = findings.get('red_flags', [])
        total_risks = (len(signals) if isinstance(signals, list) else 0) + (len(red_flags) if isinstance(red_flags, list) else 0)
        from app.services.audit_engine import SCORING_FUNCTIONS
        actual = SCORING_FUNCTIONS[func_name](findings)
        formula = f'基准分{base}分 - (风险信号数 × 信号权重)'
        if total_risks > 0:
            deductions.append({'reason': f'检测到{total_risks}个风险信号', 'amount': round(base - actual, 1)})

    elif func_name == 'ai_doc_review':
        axes = findings.get('axes', {})
        from app.services.audit_engine import SCORING_FUNCTIONS
        actual = SCORING_FUNCTIONS[func_name](findings)
        formula = f'基准分{base}分，各维度得分(0-10)平均值 × 10'
        if axes:
            axis_scores = []
            for name, ad in axes.items():
                s = ad.get('score', 5) if isinstance(ad, dict) else 5
                axis_scores.append(f'{name}:{s}/10')
            detail = '维度得分: ' + ', '.join(axis_scores[:5])
            if actual < base:
                deductions.append({'reason': f'部分维度得分偏低', 'amount': round(base - actual, 1)})

    elif func_name == 'style_analysis':
        from app.services.audit_engine import SCORING_FUNCTIONS
        actual = SCORING_FUNCTIONS[func_name](findings)
        form = findings.get('formality_level', 50)
        formula = f'(正式度{form} + 一致性) / 2'
        if actual < base:
            deductions.append({'reason': f'正式度或一致性偏低', 'amount': round(base - actual, 1)})

    else:
        from app.services.audit_engine import SCORING_FUNCTIONS
        actual = SCORING_FUNCTIONS.get(func_name, lambda f: 50)(findings)
        formula = '未定义的评分函数'

    return {
        'base_value': base,
        'actual_score': round(actual, 1),
        'deductions': deductions,
        'formula_zh': formula,
        'detail_zh': detail,
    }


def _ensure_dir():
    os.makedirs(str(AUDIT_REPORTS_DIR), exist_ok=True)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX Generation
# ═══════════════════════════════════════════════════════════════════════════════

def generate_docx(run_id: int, folder_ids: list[int],
                  enabled_functions: list[str]) -> str | None:
    """Generate a scored DOCX report with table of contents."""
    _ensure_dir()
    data = _load_run_data(run_id)
    if not data:
        logger.error(f"No data for run {run_id}")
        return None

    doc = Document()

    # --- Professional Chinese formatting ---
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.font.name = '宋体'
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    for level in range(1, 5):
        hs_name = f'Heading {level}'
        if hs_name in [s.name for s in doc.styles]:
            hs = doc.styles[hs_name]
            hs.font.name = '黑体'
            hs.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # --- Title ---
    title = doc.add_heading('投标文件审计报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"生成时间: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}  |  "
        f"运行ID: {run_id}  |  "
        f"状态: {'未通过' if data.get('overall_status') == 'FAIL' else '通过'}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- 1. 执行摘要 ---
    doc.add_heading('1. 执行摘要', level=1)
    overall_score = data.get('overall_score', 0) or 0
    overall_status = data.get('overall_status', 'N/A')

    summary = doc.add_paragraph()
    summary.add_run('综合得分: ').bold = True
    score_run = summary.add_run(f'{overall_score:.1f}/100')
    score_run.bold = True
    score_run.font.color.rgb = RGBColor(220, 38, 38) if overall_status == 'FAIL' else RGBColor(34, 197, 94)

    summary2 = doc.add_paragraph()
    summary2.add_run('综合状态: ').bold = True
    status_text = '未通过 (FAIL)' if overall_status == 'FAIL' else '通过 (PASS)'
    status_run = summary2.add_run(status_text)
    status_run.bold = True
    status_run.font.color.rgb = RGBColor(220, 38, 38) if overall_status == 'FAIL' else RGBColor(34, 197, 94)

    doc.add_paragraph(f"投标人数量: {data.get('bidder_count', 0)}  |  文件数量: {data.get('file_count', 0)}")

    # Top risks
    risks = _collect_top_risks(data)
    if risks:
        doc.add_heading('主要风险', level=2)
        for i, risk in enumerate(risks[:5], 1):
            doc.add_paragraph(f"{i}. [{risk['bidder']}] {risk['function']}: {risk['detail']}", style='List Number')

    # Fatal-flaw note
    if overall_status == 'FAIL':
        p = doc.add_paragraph()
        p.add_run('触发未通过阈值 (FATAL-FLAW): ').bold = True
        p.runs[0].font.color.rgb = RGBColor(220, 38, 38)
        p.add_run('一个或多个审计功能的得分低于其未通过阈值。详见下方各功能详细分析。')

    doc.add_page_break()

    # --- 2. 投标人对比 ---
    doc.add_heading('2. 投标人对比', level=1)
    comparison = _build_comparison_matrix(data, enabled_functions)
    if comparison['bidders']:
        _write_comparison_table(doc, comparison, enabled_functions)
    else:
        doc.add_paragraph('（仅有一个投标人，无法进行对比）')

    doc.add_page_break()

    # --- 3. 各投标人详细分析 ---
    doc.add_heading('3. 各投标人详细分析', level=1)
    bidders = _group_by_bidder(data)
    for idx, (bidder_name, bidder_data) in enumerate(bidders.items(), 1):
        bidder_scores = []
        for fname, results in bidder_data['files'].items():
            for r in results:
                bidder_scores.append(r.get('score', 0) or 0)
        bidder_avg = sum(bidder_scores) / len(bidder_scores) if bidder_scores else 0
        color = RGBColor(220, 38, 38) if bidder_avg < 50 else (RGBColor(245, 158, 11) if bidder_avg < 70 else RGBColor(34, 197, 94))
        h = doc.add_heading(f'3.{idx} {bidder_name} — 综合得分: {bidder_avg:.1f}', level=2)
        for run in h.runs:
            run.font.color.rgb = color
        _write_bidder_detail(doc, bidder_name, bidder_data, enabled_functions)
        doc.add_page_break()

    # --- 4. 附录 ---
    doc.add_heading('4. 附录', level=1)
    doc.add_heading('4.1 审计配置', level=2)
    config = data.get('config_snapshot', {}) or {}
    cfg_table = doc.add_table(rows=1, cols=4)
    cfg_table.style = 'Light Grid Accent 1'
    for i, text in enumerate(['审计功能', '未通过阈值', '权重', '已启用']):
        cfg_table.rows[0].cells[i].text = text
        for p in cfg_table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for func_name, cfg in config.items():
        row = cfg_table.add_row()
        row.cells[0].text = FUNCTION_LABELS_ZH.get(func_name, func_name)
        row.cells[1].text = str(cfg.get('fail_threshold', 50))
        row.cells[2].text = f"{cfg.get('weight', 0):.3f}"
        row.cells[3].text = '是' if cfg.get('enabled', True) else '否'

    doc.add_heading('4.2 审计元数据', level=2)
    doc.add_paragraph(f"运行ID: {run_id}")
    doc.add_paragraph(f"开始时间: {data.get('started_at', 'N/A')}")
    doc.add_paragraph(f"完成时间: {data.get('completed_at', 'N/A')}")
    doc.add_paragraph(f"已启用功能: {'、'.join(FUNCTION_LABELS_ZH.get(f, f) for f in enabled_functions)}")

    # Save
    filename = f"audit_{run_id}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.docx"
    path = str(AUDIT_REPORTS_DIR / filename)
    doc.save(path)
    logger.info(f"DOCX saved: {path}")
    return path


def _collect_top_risks(data: dict) -> list[dict]:
    risks = []
    for fr in data.get('file_results', []):
        score = fr.get('score', 100) or 100
        if score < 50:
            risks.append({
                'bidder': fr.get('bidder_label', 'Unknown'),
                'function': FUNCTION_LABELS_ZH.get(fr.get('function_name', ''), fr.get('function_name', '')),
                'score': score,
                'detail': _risk_detail(fr),
            })
    risks.sort(key=lambda r: r['score'])
    return risks


def _risk_detail(fr: dict) -> str:
    findings = fr.get('findings', {}) or {}
    status = fr.get('status', '')
    if status == 'error':
        return fr.get('error_message', 'Function error')[:200]
    if status == 'skipped':
        return 'File skipped — no text content'
    func_name = fr.get('function_name', '')
    if func_name == 'typo_detection':
        count = findings.get('total_findings', 0)
        return f'{count} typo(s) found'
    if func_name == 'compliance_check':
        summary = findings.get('summary', {})
        critical = summary.get('critical', 0)
        violations = summary.get('violation', 0)
        return f'Critical: {critical}, Violations: {violations}'
    if func_name == 'quote_anomaly':
        flags = findings.get('flags', [])
        return f'{len(flags)} anomaly flag(s): {", ".join(str(f) for f in flags[:3])}'
    if func_name == 'relationship_extraction':
        signals = findings.get('collusion_signals', [])
        red_flags = findings.get('red_flags', [])
        return f'{len(signals)} signal(s), {len(red_flags)} red flag(s)'
    if func_name == 'ai_doc_review':
        if findings.get('parse_error'):
            return 'AI response could not be parsed'
        return f"Overall score from AI review"
    return f'Score: {fr.get("score", "N/A")}'


def _build_comparison_matrix(data: dict, enabled_functions: list[str]) -> dict:
    bidders = {}
    for fr in data.get('file_results', []):
        bidder = fr.get('bidder_label', 'Unknown')
        func = fr.get('function_name', '')
        score = fr.get('score', 0) or 0
        if bidder not in bidders:
            bidders[bidder] = {}
        if func not in bidders[bidder]:
            bidders[bidder][func] = []
        bidders[bidder][func].append(score)

    # Average per bidder per function
    matrix = {}
    for bidder, funcs in bidders.items():
        matrix[bidder] = {}
        for func, scores in funcs.items():
            matrix[bidder][func] = sum(scores) / len(scores) if scores else 0

    # Rank bidders by average across all functions
    ranked = []
    for bidder, funcs in matrix.items():
        if funcs:
            avg = sum(funcs.values()) / len(funcs)
            ranked.append((bidder, avg))
    ranked.sort(key=lambda x: x[1], reverse=True)

    return {'bidders': matrix, 'ranked': ranked}


def _write_comparison_table(doc, comparison: dict, enabled_functions: list[str]):
    bidders = comparison['bidders']
    ranked = comparison['ranked']

    # Rank table
    table = doc.add_table(rows=1, cols=len(enabled_functions) + 2)
    table.style = 'Light Grid Accent 1'

    hdr = table.rows[0].cells
    hdr[0].text = '排名'
    hdr[1].text = '投标人'
    for i, func in enumerate(enabled_functions):
        hdr[i + 2].text = FUNCTION_LABELS_ZH.get(func, func)
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    for rank, (bidder_name, avg_score) in enumerate(ranked, 1):
        row = table.add_row()
        cells = row.cells
        cells[0].text = str(rank)
        cells[1].text = bidder_name
        funcs = bidders.get(bidder_name, {})
        for i, func in enumerate(enabled_functions):
            score = funcs.get(func, 'N/A')
            if isinstance(score, (int, float)):
                cells[i + 2].text = f"{score:.1f}"
            else:
                cells[i + 2].text = str(score)

    doc.add_paragraph()


def _group_by_bidder(data: dict) -> dict:
    bidders = {}
    for fr in data.get('file_results', []):
        bidder = fr.get('bidder_label', 'Unknown')
        if bidder not in bidders:
            bidders[bidder] = {'files': {}, 'scores': {}}
        filename = fr.get('filename', 'unknown')
        if filename not in bidders[bidder]['files']:
            bidders[bidder]['files'][filename] = []
        bidders[bidder]['files'][filename].append(fr)
    return bidders


def _write_bidder_detail(doc, bidder_name: str, bidder_data: dict,
                         enabled_functions: list[str]):
    # Per-file score summary
    doc.add_heading('文件评分总览', level=3)
    table = doc.add_table(rows=1, cols=len(enabled_functions) + 2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '文件'
    hdr[1].text = '综合'
    for i, func in enumerate(enabled_functions):
        hdr[i + 2].text = FUNCTION_LABELS_ZH.get(func, func)
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    for filename, results in bidder_data['files'].items():
        row = table.add_row()
        cells = row.cells
        cells[0].text = filename[:40]
        scores = [r.get('score', 0) or 0 for r in results]
        avg = sum(scores) / len(scores) if scores else 0
        cells[1].text = f"{avg:.1f}"
        for i, func in enumerate(enabled_functions):
            func_results = [r for r in results if r.get('function_name') == func]
            if func_results:
                score = func_results[0].get('score', 'N/A')
                status = func_results[0].get('status', '')
                if status == 'error':
                    cells[i + 2].text = 'ERR'
                elif status == 'skipped':
                    cells[i + 2].text = 'SKIP'
                elif isinstance(score, (int, float)):
                    cells[i + 2].text = f"{score:.0f}"
                else:
                    cells[i + 2].text = str(score)
            else:
                cells[i + 2].text = '-'

    doc.add_paragraph()

    # Per-file scoring detail with rationale
    for filename, results in bidder_data['files'].items():
        doc.add_heading(f"文件: {filename[:60]}", level=3)
        for func in enabled_functions:
            func_results = [r for r in results if r.get('function_name') == func]
            if not func_results:
                p = doc.add_paragraph()
                p.add_run(f"{FUNCTION_LABELS_ZH.get(func, func)}: ").bold = True
                p.add_run('本次审计未启用此功能')
                continue

            fr = func_results[0]
            status = fr.get('status', '')
            score = fr.get('score', 'N/A')
            findings = fr.get('findings', {}) or {}

            # Score with rationale
            doc.add_heading(f"{FUNCTION_LABELS_ZH.get(func, func)}", level=4)
            p = doc.add_paragraph()
            p.add_run('得分: ').bold = True
            if status == 'error':
                score_run = p.add_run(f"错误 — {fr.get('error_message', '')[:100]}")
                score_run.font.color.rgb = RGBColor(220, 38, 38)
            elif status == 'skipped':
                p.add_run('跳过 — 文件无文本内容')
            else:
                score_val = f"{score:.0f}/100" if isinstance(score, (int, float)) else str(score)
                score_run = p.add_run(score_val)
                if isinstance(score, (int, float)):
                    score_run.font.color.rgb = RGBColor(220, 38, 38) if score < 50 else (RGBColor(245, 158, 11) if score < 70 else RGBColor(34, 197, 94))

                # Scoring rationale
                rationale = _get_scoring_rationale(func, findings)
                if rationale:
                    p2 = doc.add_paragraph()
                    p2.add_run('评分依据: ').bold = True
                    p2.add_run(f"基准分 {rationale['base_value']:.0f} 分，计算方式: {rationale['formula_zh']}")
                    if rationale['deductions']:
                        for d in rationale['deductions']:
                            p2.add_run(f"\n  └ {d['reason']}: -{d['amount']:.1f} 分")
                    if rationale['detail_zh']:
                        p2.add_run(f"\n  └ {rationale['detail_zh']}")

            # Detailed findings
            _write_function_findings(doc, func, findings)


def _write_function_findings(doc, func_name: str, findings: dict):
    if func_name == 'rule_extraction':
        rules = findings.get('rules', [])
        if rules:
            doc.add_paragraph(f"  Extracted {len(rules)} rules:")
            for rule in rules[:10]:
                if isinstance(rule, dict):
                    doc.add_paragraph(
                        f"    - [{rule.get('category', '')}] {rule.get('description', str(rule))[:150]}",
                        style='List Bullet')

    elif func_name == 'compliance_check':
        results = findings.get('results', [])
        summary = findings.get('summary', {})
        if summary:
            doc.add_paragraph(
                f"  Pass: {summary.get('pass', 0)}, Warning: {summary.get('warning', 0)}, "
                f"Violation: {summary.get('violation', 0)}, Critical: {summary.get('critical', 0)}"
            )
        for r in results[:5]:
            if isinstance(r, dict) and r.get('verdict') not in ('pass', None):
                doc.add_paragraph(
                    f"    - [{r.get('verdict', '')}] {r.get('description', str(r))[:150]}",
                    style='List Bullet')

    elif func_name == 'typo_detection':
        findings_list = findings.get('findings', [])
        if findings_list:
            doc.add_paragraph(f"  Found {len(findings_list)} typo(s):")
            for f in findings_list[:10]:
                if isinstance(f, dict):
                    doc.add_paragraph(
                        f"    - {f.get('original', '')} -> {f.get('corrected', '')} "
                        f"[{f.get('type', '')}] {f.get('context', '')[:50]}",
                        style='List Bullet')

    elif func_name == 'quote_anomaly':
        flags = findings.get('flags', [])
        severity = findings.get('severity_index', 'N/A')
        doc.add_paragraph(f"  Severity index: {severity}")
        if flags:
            doc.add_paragraph(f"  Flags: {', '.join(str(f) for f in flags[:10])}")

    elif func_name == 'relationship_extraction':
        companies = findings.get('companies', [])
        personnel = findings.get('personnel', [])
        signals = findings.get('collusion_signals', [])
        red_flags = findings.get('red_flags', [])
        doc.add_paragraph(
            f"  Companies: {len(companies)}, Personnel: {len(personnel)}, "
            f"Signals: {len(signals)}, Red flags: {len(red_flags)}"
        )
        if signals:
            for s in signals[:5]:
                doc.add_paragraph(f"    - {str(s)[:200]}", style='List Bullet')

    elif func_name == 'ai_doc_review':
        if findings.get('parse_error'):
            doc.add_paragraph(f"  (AI response could not be parsed — see raw analysis)")
        else:
            axes = findings.get('axes', {})
            if axes:
                for axis_name, axis_data in axes.items():
                    if isinstance(axis_data, dict):
                        s = axis_data.get('score', 'N/A')
                        doc.add_paragraph(f"  {axis_name}: {s}/10", style='List Bullet')
            else:
                doc.add_paragraph(f"  See raw analysis in XLSX report")

    elif func_name == 'style_analysis':
        doc.add_paragraph(
            f"  Formality: {findings.get('formality_label', 'N/A')} "
            f"({findings.get('formality_level', 'N/A')})"
        )


# ═══════════════════════════════════════════════════════════════════════════════
# XLSX Generation
# ═══════════════════════════════════════════════════════════════════════════════

def generate_xlsx(run_id: int, folder_ids: list[int],
                  enabled_functions: list[str]) -> str | None:
    """Generate a structured XLSX report with per-function sheets."""
    _ensure_dir()
    data = _load_run_data(run_id)
    if not data:
        logger.error(f"No data for run {run_id}")
        return None

    wb = Workbook()

    # --- Styles ---
    header_font = Font(bold=True, size=11)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font_white = Font(bold=True, size=11, color="FFFFFF")
    fail_fill = PatternFill(start_color="FECACA", end_color="FECACA", fill_type="solid")
    pass_fill = PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid")
    warn_fill = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    def style_header(ws, row, col_count):
        for col in range(1, col_count + 1):
            cell = ws.cell(row=row, column=col)
            cell.font = header_font_white
            cell.fill = header_fill
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center')

    def auto_width(ws, min_width=8, max_width=40):
        for col in ws.columns:
            max_len = 0
            col_letter = get_column_letter(col[0].column)
            for cell in col:
                if cell.value:
                    max_len = max(max_len, len(str(cell.value)))
            ws.column_dimensions[col_letter].width = min(max(max_len + 2, min_width), max_width)

    # --- Sheet 1: Summary ---
    ws = wb.active
    ws.title = '摘要'
    ws.append(['投标文件审计报告', f'运行ID: {run_id}'])
    ws.append(['综合得分', data.get('overall_score', 'N/A')])
    ws.append(['综合状态', '未通过' if data.get('overall_status') == 'FAIL' else '通过'])
    ws.append(['投标人数量', data.get('bidder_count', 0)])
    ws.append(['文件数量', data.get('file_count', 0)])
    ws.append(['开始时间', data.get('started_at', 'N/A')])
    ws.append(['完成时间', data.get('completed_at', 'N/A')])
    ws.append([])

    # Score matrix
    comparison = _build_comparison_matrix(data, enabled_functions)
    headers = ['投标人'] + [FUNCTION_LABELS_ZH.get(f, f) for f in enabled_functions] + ['综合']
    ws.append(headers)
    style_header(ws, ws.max_row, len(headers))

    ranked = comparison.get('ranked', [])
    for bidder_name, avg_score in ranked:
        funcs = comparison['bidders'].get(bidder_name, {})
        row = [bidder_name]
        for func in enabled_functions:
            score = funcs.get(func, 'N/A')
            row.append(round(score, 1) if isinstance(score, (int, float)) else score)
        row.append(round(avg_score, 1))
        ws.append(row)

    auto_width(ws)

    # --- Per-function sheets ---
    results_by_func = {}
    for fr in data.get('file_results', []):
        func = fr.get('function_name', '')
        if func not in results_by_func:
            results_by_func[func] = []
        results_by_func[func].append(fr)

    for func_name in enabled_functions:
        sheet_name = FUNCTION_LABELS_ZH.get(func_name, func_name)[:31]
        ws_func = wb.create_sheet(title=sheet_name)

        func_results = results_by_func.get(func_name, [])
        headers = ['投标人', '文件', '得分', '状态', '重试', '错误信息', '详细发现']
        ws_func.append(headers)
        style_header(ws_func, 1, len(headers))

        for fr in func_results:
            findings_summary = _summarize_findings(func_name, fr.get('findings', {}) or {})
            ws_func.append([
                fr.get('bidder_label', ''),
                fr.get('filename', ''),
                fr.get('score', 'N/A'),
                fr.get('status', ''),
                fr.get('retry_count', 0),
                (fr.get('error_message', '') or '')[:200],
                findings_summary,
            ])

        auto_width(ws_func)

        # Color rows
        for row in range(2, ws_func.max_row + 1):
            score_cell = ws_func.cell(row=row, column=3)
            if isinstance(score_cell.value, (int, float)):
                if score_cell.value < 50:
                    for col in range(1, len(headers) + 1):
                        ws_func.cell(row=row, column=col).fill = fail_fill
                elif score_cell.value < 70:
                    for col in range(1, len(headers) + 1):
                        ws_func.cell(row=row, column=col).fill = warn_fill
                else:
                    for col in range(1, len(headers) + 1):
                        ws_func.cell(row=row, column=col).fill = pass_fill

    # --- Comparison Sheet ---
    if len(ranked) >= 2:
        ws_comp = wb.create_sheet(title='对比')
        headers = ['排名', '投标人'] + [FUNCTION_LABELS_ZH.get(f, f) for f in enabled_functions] + ['综合']
        ws_comp.append(headers)
        style_header(ws_comp, 1, len(headers))

        for rank, (bidder_name, avg_score) in enumerate(ranked, 1):
            funcs = comparison['bidders'].get(bidder_name, {})
            row = [rank, bidder_name]
            for func in enabled_functions:
                score = funcs.get(func, 'N/A')
                row.append(round(score, 1) if isinstance(score, (int, float)) else score)
            row.append(round(avg_score, 1))
            ws_comp.append(row)

        auto_width(ws_comp)

    # --- Config Sheet ---
    ws_cfg = wb.create_sheet(title='配置')
    headers = ['审计功能', '已启用', '未通过阈值', '权重', '严重度阈值']
    ws_cfg.append(headers)
    style_header(ws_cfg, 1, len(headers))

    config = data.get('config_snapshot', {}) or {}
    for func_name, cfg in config.items():
        ws_cfg.append([
            FUNCTION_LABELS_ZH.get(func_name, func_name),
            cfg.get('enabled', True),
            cfg.get('fail_threshold', 50),
            f"{cfg.get('weight', 0):.3f}",
            json.dumps(cfg.get('severity_thresholds', {}), ensure_ascii=False),
        ])
    auto_width(ws_cfg)

    # Save
    filename = f"audit_{run_id}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.xlsx"
    path = str(AUDIT_REPORTS_DIR / filename)
    wb.save(path)
    logger.info(f"XLSX saved: {path}")
    return path


def _summarize_findings(func_name: str, findings: dict) -> str:
    if not findings:
        return ''
    if func_name == 'rule_extraction':
        return f"提取{len(findings.get('rules', []))}条规则"
    if func_name == 'compliance_check':
        s = findings.get('summary', {})
        return f"通过:{s.get('pass',0)} 警告:{s.get('warning',0)} 违规:{s.get('violation',0)} 严重:{s.get('critical',0)}"
    if func_name == 'typo_detection':
        return f"发现{findings.get('total_findings', 0)}个错别字 ({', '.join(findings.get('layers_run', []))})"
    if func_name == 'quote_anomaly':
        return f"严重度:{findings.get('severity_index', '?')}, 异常标志:{len(findings.get('flags', []))}"
    if func_name == 'relationship_extraction':
        return f"公司:{len(findings.get('companies',[]))}, 风险信号:{len(findings.get('collusion_signals',[]))}"
    if func_name == 'ai_doc_review':
        if findings.get('parse_error'):
            return '解析错误'
        axes = findings.get('axes', {})
        if axes:
            return ', '.join(f"{k}:{v.get('score','?') if isinstance(v, dict) else v}" for k, v in list(axes.items())[:3])
        return '见原始数据'
    if func_name == 'style_analysis':
        return f"正式度: {findings.get('formality_level', '?')}"
    return json.dumps(findings, ensure_ascii=False)[:200]