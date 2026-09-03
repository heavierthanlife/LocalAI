"""Blueprint: credit routes (auto-extracted)."""
import os, json, uuid, time, logging, hashlib, io, threading
from functools import wraps
from flask import Blueprint, request, jsonify, session, send_file, render_template, url_for, current_app

from app.config import BASE_DIR, DATA_DIR, TEMP_ROOT, TEMP_DIR, USER_FILES_ORIGINAL_ROOT, CREDIT_REPORTS_DIR, PROJECT_FILES_ROOT, to_rel_path, resolve_path
from app.database import get_db_connection, db_transaction
from app.utils.helpers import utc_now, beijing_now, safe_error_response, split_thinking_answer
from app.services.credit_task_registry import (
    set_task as _credit_set, get_task as _credit_get,
    patch_task as _credit_patch, task_exists as _credit_exists,
    delete_task as _credit_delete,
)
from app.services.file_cache import file_cache_manager, add_to_cache, load_cache_from_db
from app.services.session_manager import get_user_id

from io import BytesIO
from psycopg2.extras import RealDictCursor
from app.services.credit_checker import CreditChecker

from docx import Document
from docx.shared import Cm, Inches

logger = logging.getLogger(__name__)

# Rate limiter for credit check endpoints (prevents abuse of browser automation)
# QA-Loop C4: previously a process-local dict — under gunicorn's multiple
# gevent workers each worker had its own counter, letting users bypass the
# 10/5min cap by spreading requests across workers. Now Redis-backed
# (credit_rate:{ip}) with in-memory fallback when Redis is unavailable,
# mirroring credit_task_registry.py's degradation pattern.
_credit_rate_limit = {}
_CREDIT_RATE_MAX = 10       # max requests per window
_CREDIT_RATE_WINDOW = 300   # 5 minutes in seconds
_RATE_PREFIX = 'credit_rate:'

def _rate_limit_check(ip):
    """Return True if allowed, False if limited. Redis-primary, memory fallback."""
    now = time.time()
    try:
        from app.services.redis_client import get_redis
        r = get_redis(decode_responses=True)
        if r is not None:
            key = _RATE_PREFIX + (ip or 'unknown')
            count = r.incr(key)
            if count == 1:
                r.expire(key, _CREDIT_RATE_WINDOW)
            return count <= _CREDIT_RATE_MAX
    except Exception as e:
        logger.warning(f"Credit rate limiter Redis failed, falling back to memory: {e}")
    # Fallback: in-memory (dev / single-worker)
    key = f'credit:{ip}'
    for k in list(_credit_rate_limit.keys()):
        if _credit_rate_limit[k]['ts'] < now - _CREDIT_RATE_WINDOW:
            del _credit_rate_limit[k]
    entry = _credit_rate_limit.get(key)
    if entry is None:
        _credit_rate_limit[key] = {'count': 0, 'ts': now}
        entry = _credit_rate_limit[key]
    elif entry['count'] >= _CREDIT_RATE_MAX:
        return False
    entry['count'] += 1
    return True

def credit_rate_limit(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        ip = request.remote_addr or 'unknown'
        if not _rate_limit_check(ip):
            logger.warning(f"Credit rate limit exceeded for {ip}")
            return jsonify({"error": "Too many credit check requests. Please try again later."}), 429
        return f(*args, **kwargs)
    return wrapper

credit_bp = Blueprint('credit', __name__, template_folder=str(BASE_DIR / 'templates'), static_folder=str(BASE_DIR / 'static'))

def _require_registered():
    """Helper: all credit routes need consent + login. Returns user_id or error Response."""
    if session.get('consent_value', 0) != 1:
        return jsonify({"error": "请先登录"}), 401
    uid = session.get('user_id')
    if not uid:
        return jsonify({"error": "请先创建账户"}), 401
    return uid

@credit_bp.route('/start_credit_check', methods=['POST'])
@credit_rate_limit
def start_credit_check():
    uid_or_err = _require_registered()
    if not isinstance(uid_or_err, str):
        return uid_or_err
    user_id = uid_or_err
    data = request.get_json()
    companies = data.get('companies', [])
    urls = data.get('urls', [])
    if not companies or not urls:
        return jsonify({"error": "Need companies and urls"}), 400

    task_id = str(uuid.uuid4())
    user_id = get_user_id()
    with current_app.app_context():
        download_url = url_for('download_credit_report', task_id=task_id, _external=True)

    _credit_set(task_id, {
        'status': 'running',
        'progress': 0,
        'total': len(companies),
        'captcha_needed': False,
        'captcha_image': None,
        'captcha_task': None,
        'captcha_solution': None,
        'reload_captcha': False,
        'download_url': download_url,
        'error': None,
        'waiting': False,
        'resume': False
    })

    threading.Thread(target=_run_credit_check,
                     args=(task_id, companies, urls, user_id, True),
                     daemon=True).start()
    return jsonify({"task_id": task_id})

def _run_credit_check(task_id, companies, urls, user_id, manual_mode=True):
    checker = CreditChecker()
    screenshots = {}

    try:
        for idx, company in enumerate(companies):
            company_shots = []
            for url in urls:
                logger.info(f"Processing {company} at {url}")

                # Navigate and fill company name, also sets zoom
                checker.navigate_and_fill(company, url)

                # Handle CAPTCHA if present
                if checker._is_captcha_present():
                    captcha_img = checker.get_captcha_element_screenshot()
                    if captcha_img:
                        _credit_patch(task_id,
                                      captcha_needed=True,
                                      captcha_image=captcha_img.getvalue(),
                                      captcha_solution=None,
                                      reload_captcha=False)

                        # Wait for user to solve CAPTCHA via modal
                        captcha_start = time.time()
                        captcha_timeout = 300  # 5 minutes max wait
                        while True:
                            if time.time() - captcha_start > captcha_timeout:
                                logger.warning(f"CAPTCHA wait timed out for task {task_id}")
                                _credit_patch(task_id, status='error', error='验证码等待超时')
                                return
                            task_state = _credit_get(task_id) or {}
                            solution = task_state.get('captcha_solution')
                            reload_flag = task_state.get('reload_captcha', False)

                            if solution is not None:
                                break
                            if reload_flag:
                                _credit_patch(task_id, reload_captcha=False)
                                # Refresh the CAPTCHA image on the page
                                checker.refresh_captcha()
                                time.sleep(1)
                                new_img = checker.get_captcha_element_screenshot()
                                if new_img:
                                    _credit_patch(task_id, captcha_image=new_img.getvalue())
                            time.sleep(1)

                        # Submit the CAPTCHA solution
                        task_state = _credit_get(task_id) or {}
                        solution = task_state.get('captcha_solution')
                        checker.submit_captcha(solution)

                        # Clear CAPTCHA flags
                        _credit_patch(task_id, captcha_needed=False, captcha_image=None)
                        time.sleep(3)   # wait for page to reload

                # Wait for user to confirm results
                if manual_mode:
                    _credit_patch(task_id, waiting=True, resume=False)

                    manual_start = time.time()
                    manual_timeout = 600  # 10 minutes max wait
                    while True:
                        if time.time() - manual_start > manual_timeout:
                            logger.warning(f"Manual confirm wait timed out for task {task_id}")
                            _credit_patch(task_id, waiting=False)
                            break
                        task_state = _credit_get(task_id) or {}
                        resume = task_state.get('resume', False)
                        if resume:
                            break
                        time.sleep(1)

                    _credit_patch(task_id, waiting=False)
                else:
                    time.sleep(3)   # fallback delay for full-auto (unused)

                # Capture screenshot
                shot = checker.capture_viewport()
                company_shots.append(shot)
                logger.info(f"Screenshot captured for {company} at {url}")

            screenshots[company] = company_shots
            _credit_patch(task_id, progress=idx + 1)

        # ========== Generate Word Document ==========
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(0.3)
            section.bottom_margin = Cm(0.3)
            section.left_margin = Cm(0.3)
            section.right_margin = Cm(0.3)

        first = True
        for company, shots in screenshots.items():
            if not first:
                doc.add_page_break()
            first = False
            doc.add_heading(company, level=1)
            for shot in shots:
                shot.seek(0)
                doc.add_picture(shot, width=Inches(7.2))
                doc.add_paragraph()

        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        # Save report file
        os.makedirs(CREDIT_REPORTS_DIR, exist_ok=True)
        file_name = f"credit_report_{task_id}.docx"
        file_path = os.path.join(CREDIT_REPORTS_DIR, file_name)
        with open(file_path, 'wb') as f:
            f.write(doc_buffer.getvalue())

        # Insert into database
        with get_db_connection() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "INSERT INTO credit_check_reports (user_id, task_id, file_path, companies_count) VALUES (%s, %s, %s, %s)",
                    (user_id, task_id, to_rel_path(file_path), len(companies))
                )
                conn.commit()

        # Mark task as completed
        _credit_patch(task_id, status='completed')
        logger.info(f"Credit check task {task_id} finished successfully")

    except Exception as e:
        logger.error(f"Credit check task {task_id} failed: {e}", exc_info=True)
        _credit_patch(task_id, status='error', error=str(e))
    finally:
        checker.close()

@credit_bp.route('/credit_check_status/<task_id>')
def credit_check_status(task_id):
    if session.get('consent_value', 0) != 1:
        return jsonify({"error": "请先登录"}), 401
    task = _credit_get(task_id)
    if not task:
        return jsonify({"error": "Task not found"}), 404
    result = {
        'status': task['status'],
        'progress': task['progress'],
        'total': task['total'],
        'captcha_needed': task.get('captcha_needed', False),
        'download_url': task.get('download_url'),
        'error': task.get('error'),
        'waiting': task.get('waiting', False)
    }
    return jsonify(result)

@credit_bp.route('/credit_check_resume/<task_id>', methods=['POST'])
def credit_check_resume(task_id):
    if session.get('consent_value', 0) != 1:
        return jsonify({"error": "请先登录"}), 401
    if _credit_exists(task_id):
        _credit_patch(task_id, resume=True)
    return jsonify({"status": "ok"})

@credit_bp.route('/get_captcha_image/<task_id>')
def get_captcha_image(task_id):
    if session.get('consent_value', 0) != 1:
        return "Not authorized", 401
    task = _credit_get(task_id)
    if not task or not task.get('captcha_needed') or not task.get('captcha_image'):
        return "No captcha image", 404
    img_bytes = task['captcha_image']
    return send_file(BytesIO(img_bytes), mimetype='image/png')

@credit_bp.route('/reload_captcha/<task_id>', methods=['POST'])
def reload_captcha(task_id):
    if session.get('consent_value', 0) != 1:
        return jsonify({"error": "请先登录"}), 401
    if _credit_exists(task_id):
        _credit_patch(task_id, reload_captcha=True)
    return jsonify({"status": "reloading"})

@credit_bp.route('/solve_captcha/<task_id>', methods=['POST'])
def solve_captcha(task_id):
    if session.get('consent_value', 0) != 1:
        return jsonify({"error": "请先登录"}), 401
    data = request.get_json()

    solution = data.get('solution', '')
    if _credit_exists(task_id):
        _credit_patch(task_id, captcha_solution=solution)
    return jsonify({"status": "ok"})

@credit_bp.route('/download_credit_report/<task_id>')
def download_credit_report(task_id):
    """Any registered user can download any completed credit report."""
    if session.get('consent_value', 0) != 1:
        return jsonify({"error": "请先登录"}), 401
    report_dir = str(CREDIT_REPORTS_DIR)
    file_path = os.path.join(report_dir, f"credit_report_{task_id}.docx")
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True, download_name=f"credit_report_{task_id}.docx")
    return "Report not found", 404

@credit_bp.route('/list_credit_reports')
def list_credit_reports():
    """Return ALL credit reports visible to all registered users."""
    if session.get('consent_value', 0) != 1:
        return jsonify({"error": "请先登录"}), 401
    with get_db_connection() as conn:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT r.id, r.task_id, r.companies_count, r.created_at, r.user_id,
                       u.username as created_by_name
                FROM credit_check_reports r
                JOIN users u ON r.user_id = u.user_id
                ORDER BY r.created_at DESC
                LIMIT 50
            """)
            reports = cur.fetchall()
    return jsonify({"reports": reports})

@credit_bp.route('/delete_credit_report/<int:report_id>', methods=['POST'])
def delete_credit_report(report_id):
    """Only admin can delete credit reports."""
    if session.get('consent_value', 0) != 1:
        return jsonify({"error": "请先登录"}), 401
    if session.get('role') != 'admin':
        return jsonify({"error": "仅管理员可删除报告"}), 403
    with get_db_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT file_path FROM credit_check_reports WHERE id = %s", (report_id,))
            row = cur.fetchone()
            if not row:
                return jsonify({"error": "Report not found"}), 404
            file_path = resolve_path(row[0])
            if os.path.exists(file_path):
                os.remove(file_path)
            cur.execute("DELETE FROM credit_check_reports WHERE id = %s", (report_id,))
            conn.commit()
    return jsonify({"success": True})

